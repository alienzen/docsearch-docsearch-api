# search_api.py — API de recherche avec filtrage ACL
# Mis à jour le 08/07/2026 — ES 9.4.2 · Tika 3.3.1.0 · ACL · multi-source

import os
ES_HOST = os.getenv("ES_HOST", "http://localhost:9200")
REDIS_HOST = os.getenv("REDIS_HOST", "redis")
LDAP_ENABLED = os.getenv("LDAP_ENABLED", "false").lower() == "true"

import subprocess
import tempfile
import logging
import io
import re
import json
import time
from pathlib import Path
from datetime import datetime, timezone

from fastapi import FastAPI, HTTPException, Depends, Request, Query
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel, Field
from elasticsearch import Elasticsearch
from elasticsearch.helpers import scan as es_scan
from auth.deps import current_user, optional_user, require_admin, is_admin
from auth.directory import get_effective_groups
from auth.router import router as auth_router
import version
import cluster_status
import admin_scan
import filetype_config
import runtime_config
import search_query
import path_filter
import search_log
import user_history
import log_retention
import duplicates
import synonyms
import pinned
import nps_log
import suggestion_log
import engagement_config
import ui_config
import saved_searches
import alert_notifications
import saved_collections
import custom_keywords
import audit_log
import file_sources_config
from file_sources_config import ES_SEARCH_ALIAS, DEFAULT_SOURCE_NAME
import sql_sources_config
import sql_dsn_registry
import web_sources_config

# Sans cet appel, le logger racine n'a AUCUN handler : uvicorn ne
# configure que ses propres loggers ("uvicorn.*", propagate=False), et
# tous les logger.info/debug des modules de l'application étaient donc
# écrits nulle part — seuls les WARNING et au-delà ressortaient, via le
# handler de dernier recours de la bibliothèque standard, sans horodatage
# ni niveau. Même format que app/alert_worker.py, pour que les deux se
# lisent pareil dans journalctl.
logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO"),
    format="%(asctime)s %(levelname)s %(message)s",
)

logger = logging.getLogger(__name__)

# La version vient du fichier VERSION du dépôt via version.py — plus de
# valeur en dur ici, qui restait figée à chaque livraison.
app = FastAPI(title="DocSearch API", version=version.VERSION)

# Routes /auth/* — connexion, session, SSO Kerberos, et les deux cibles
# internes du auth_request de Nginx (/auth/check-access, /auth/check-admin).
# Voir app/auth/router.py.
app.include_router(auth_router)

ES_HOST = ES_HOST
es = Elasticsearch(ES_HOST, retry_on_timeout=True, max_retries=3, request_timeout=60)

# Nombre de termes que chaque shard remonte pour une agrégation de
# facette — voir facet_agg() dans search(), qui explique pourquoi le
# défaut d'Elasticsearch fausse les comptes affichés.
FACET_SHARD_SIZE = 500

# Au-delà de cette durée totale, une recherche laisse une ligne WARNING
# dans le journal (journalctl -u docsearch-api). Une ligne par recherche
# noierait le journal sans rien apprendre : seules les lentes méritent
# d'être signalées, les autres restent mesurées et enregistrées dans
# l'index search_logs.
#
# 2000 ms n'est pas une valeur inventée ici : c'est déjà le seuil de
# déclenchement de la supervision ({$DOCSEARCH.RECHERCHE.MS.MAX}, voir
# docsearch-infra/zabbix/REFERENCE.md). Les deux doivent rester alignés,
# sinon Zabbix alerte sur des recherches dont le journal ne dit rien —
# ou l'inverse.
#
# 0 désactive complètement la ligne de journal (la mesure, elle,
# continue).
SLOW_SEARCH_MS = int(os.getenv("SLOW_SEARCH_MS", "2000"))


@app.middleware("http")
async def audit_log_middleware(request: Request, call_next):
    """
    Journal d'audit générique des actions d'administration — voir
    audit_log.py. Volontairement un middleware plutôt qu'un appel
    explicite dans chaque route /admin/* : une nouvelle route de mutation
    est ainsi auditée automatiquement dès sa création, sans modification
    de ce fichier ni oubli possible.

    Ne journalise que les mutations (POST/DELETE/PUT) sous /admin/* dont
    la réponse est un succès (< 400) — un échec (validation, 404, Redis/ES
    injoignable...) ne représente aucun changement réel, l'enregistrer
    serait trompeur. Le corps de la requête est lu ICI, avant call_next :
    Starlette met en cache les octets déjà lus (request._body), la route
    elle-même peut donc ensuite reconstruire son modèle Pydantic à partir
    du corps sans rien perdre.
    """
    is_mutation = (
        request.method in ("POST", "DELETE", "PUT")
        and request.url.path.startswith("/admin/")
    )
    body_bytes = await request.body() if is_mutation else b""

    response = await call_next(request)

    if is_mutation and response.status_code < 400:
        # scope["route"] n'est renseigné qu'après résolution du routage,
        # qui a lieu à l'intérieur de call_next — d'où sa lecture ici,
        # après coup. .path est le PATRON de route (ex:
        # "/admin/file-sources/{name}/label"), request.path_params les valeurs
        # résolues (ex: {"name": "finance"}) — les deux ensemble
        # permettent de reconstituer une action lisible côté UI sans
        # dépendre d'une regex sur l'URL brute.
        route = request.scope.get("route")
        path_template = getattr(route, "path", request.url.path)
        try:
            body = json.loads(body_bytes) if body_bytes else {}
        except json.JSONDecodeError:
            body = {}
        audit_log.log_action(
            es,
            # La route a déjà exigé require_admin : l'identité est celle du
            # jeton, jamais celle d'un en-tête. Le middleware ne peut pas
            # utiliser Depends (il tourne hors du routage), d'où cette
            # résolution directe — même code, même vérification.
            username=optional_user(request, request.headers.get("x-user")) or "inconnu",
            method=request.method,
            path=path_template,
            path_params=dict(request.path_params),
            body=body if isinstance(body, dict) else {},
            status_code=response.status_code,
        )
    return response


# ── Santé ────────────────────────────────────────────────────
@app.get("/health")
def health():
    try:
        info = es.info()
        return {
            "status":       "ok",
            # Version de DocSearch, à ne pas confondre avec "es_version"
            # juste en dessous. Exposée ici parce que c'est la route que
            # l'exploitation interroge au curl, sans session — même
            # niveau de divulgation que le nom du cluster et la version
            # d'Elasticsearch, déjà publics sur cette route.
            "version":      version.VERSION,
            "commit":       version.COMMIT,
            "build_date":   version.BUILD_DATE,
            "es_version":   info["version"]["number"],
            "cluster":      info["cluster_name"],
            "acl_enabled":  True,
            "ldap_enabled": str(LDAP_ENABLED),
        }
    except Exception as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


# ── Modèle de requête ────────────────────────────────────────
class SearchQuery(BaseModel):
    query:           str
    size:            int = 10
    from_:           int = Field(0, alias="from")
    sort:            str = "_score"
    extension:       str | list[str] | None = None
    has_attachments: bool | None = None
    date_from:       str | None = None   # filtre sur date_modified (voir build de la requête)
    date_to:         str | None = None   # idem
    author:          str | list[str] | None = None
    folder:          str | list[str] | None = None   # sélection cumulative, comme extension/author/source
    keywords:        str | list[str] | None = None   # sélection cumulative, comme extension/author/folder/source
    source:          str | list[str] | None = None   # nom(s) de source (file_sources_config.py) — absent = recherche fédérée sur toutes
    search_in:       str = "all"   # "all" | "title" | "author" | "keywords" | "filepath" — restreint le champ interrogé
    # Recherche exacte : les mots sont cherchés tels qu'écrits, sans
    # racinisation ni synonymes ni tolérance aux fautes — mais toujours
    # aux accents et à la casse près (« Congrès » = « CONGRES »). Se
    # combine avec les guillemets, qui restent la façon d'exiger en plus
    # l'ordre et l'adjacence des mots (voir build_text_clause).
    exact:           bool = False
    # Facettes personnalisées par source SQL (voir sql_sources_config.py:
    # FieldMapping.facet) — {es_field: [valeurs sélectionnées]}, sélection
    # cumulative comme les autres facettes. Les clés qui ne correspondent
    # à aucune facette active de la/des source(s) en jeu sont ignorées
    # (voir _active_custom_facets()).
    custom:          dict[str, list[str]] | None = None

    model_config = {"populate_by_name": True}


class SavedSearchCreate(BaseModel):
    # Reflète directement l'état de l'UI (voir `state` dans index.html),
    # pas les valeurs résolues envoyées à /search (ex: "ext" est la ou
    # les clés de chip sélectionnées — "word" — pas la liste
    # d'extensions qu'elles recouvrent — [docx, doc]) : ça permet de
    # restaurer l'interface (chips actifs, champs) directement depuis
    # l'enregistrement, sans avoir à inverser une résolution.
    name:      str
    query:     str
    search_in: str = "all"
    # Rejoué tel quel par la vérification d'alertes (voir
    # build_query_clauses) : une recherche exacte enregistrée avec une
    # alerte doit notifier sur les mêmes documents qu'elle affiche, et
    # oublier ce critère ici la ferait notifier bien plus large.
    exact:     bool = False
    ext:       str | list[str] = "all"
    author:    str | list[str] | None = None
    folder:    str | list[str] | None = None
    keywords:  str | list[str] | None = None
    source:    str | list[str] | None = None
    custom:    dict[str, list[str]] | None = None
    date_from: str | None = None
    date_to:   str | None = None
    sort:      str = "_score"


class SavedSearchAlertUpdate(BaseModel):
    enabled:   bool
    frequency: str = "daily"   # "daily" | "weekly" — voir saved_searches.set_alert()


class SavedCollectionCreate(BaseModel):
    name: str


class SavedCollectionRename(BaseModel):
    name: str


class SavedCollectionDocumentAdd(BaseModel):
    doc_id: str


class SavedCollectionSharing(BaseModel):
    groups: list[str] = []


# ── Filtre ACL ───────────────────────────────────────────────
def build_acl_filter(username: str) -> dict:
    """
    Filtre Elasticsearch garantissant qu'un utilisateur
    ne voit que les documents auxquels il a accès :
      - documents publics (acl.public = true)
      - documents dont il est propriétaire
      - documents partagés explicitement avec lui
      - documents partagés avec un de ses groupes (POSIX ou AD)
    """
    user_groups = get_effective_groups(username)

    return {
        "bool": {
            "should": [
                {"term":  {"acl.public": True}},
                {"term":  {"acl.owner":  username}},
                {"term":  {"acl.users":  username}},
                {"terms": {"acl.groups": user_groups}} if user_groups
                else {"term": {"acl.groups": "__never__"}},
            ],
            "minimum_should_match": 1,
        }
    }


def _folder_filter(folder: str | list[str] | None) -> dict | None:
    """
    Filtre ES pour la facette "Dossier" — sélection cumulative (comme
    extension/author/source) : matche tout document sous N'IMPORTE LEQUEL
    des dossiers demandés, exact OU sous-dossier (ex: folder="Finance"
    matche "Finance" et "Finance/Rapports"). Chaque dossier ajoute sa
    propre paire term/prefix au should, combinées en OR.
    """
    if not folder:
        return None
    folders = folder if isinstance(folder, list) else [folder]
    should = []
    for f in folders:
        should.append({"term": {"folder": f}})
        should.append({"prefix": {"folder": f.rstrip("/") + "/"}})
    return {"bool": {"should": should, "minimum_should_match": 1}}


def _keywords_filter(keywords: str | list[str] | None) -> dict | None:
    """
    Filtre ES pour la facette "Mots-clés" — sélection cumulative en ET :
    un document doit porter TOUS les mots-clés demandés, donc chaque clic
    RESTREINT le résultat. C'est la seule facette dans ce cas : un
    `{"terms": {...}}` unique, utilisé partout ailleurs, est un OU, et
    cocher un second mot-clé ajoutait des résultats au lieu d'en retirer.

    Le ET n'a de sens ici que parce que `keywords` est multi-valué
    (indexer.py: get_keywords() renvoie une liste). extension/author/
    folder/source et les facettes SQL ne portent qu'une valeur par
    document — un ET n'y matcherait jamais rien, d'où le OU conservé.

    Toute évolution ici est à répercuter dans search_query.py, qui en
    tient une copie pour le worker d'alertes (voir son en-tête).
    """
    if not keywords:
        return None
    kws = keywords if isinstance(keywords, list) else [keywords]
    return {"bool": {"filter": [{"term": {"keywords": k}} for k in kws]}}



def get_client_ip(request: Request) -> str | None:
    """
    Résout l'IP réelle du client. En production, Nginx est devant l'API
    (voir nginx.conf) et transmet X-Forwarded-For / X-Real-IP — sans ça,
    request.client.host ne serait que l'IP interne de Nginx, pas celle
    de l'utilisateur. X-Forwarded-For peut contenir une chaîne de proxies
    ("client, proxy1, proxy2") : le premier maillon est le client d'origine.
    """
    xff = request.headers.get("x-forwarded-for")
    if xff:
        return xff.split(",")[0].strip()
    real_ip = request.headers.get("x-real-ip")
    if real_ip:
        return real_ip
    return request.client.host if request.client else None


def _ensure_index_exists():
    """
    Vérifie que l'alias fédéré (ES_SEARCH_ALIAS) existe avant toute
    requête qui le suppose — sans ça, une installation fraîche (avant le
    tout premier ./manage.sh init) remonte une exception ES non gérée,
    traduite par FastAPI en 500 générique ('Internal Server Error') sans
    aucune indication utile. L'alias est créé par create_index() (voir
    docsearch-ingestion/indexer.py) dès la première source indexée —
    toutes les sources y contribuent, jamais un index nommé en dur.
    """
    if not es.indices.exists_alias(name=ES_SEARCH_ALIAS):
        raise HTTPException(
            status_code=503,
            detail=(
                f"Aucune source n'a encore été indexée (alias "
                f"'{ES_SEARCH_ALIAS}' introuvable). Exécutez "
                f"'./manage.sh init' depuis docsearch-infra pour indexer "
                f"la source par défaut."
            ),
        )


def _get_any_source(name: str):
    """Cherche `name` dans les trois registres (fichiers, SQL, web),
    dans cet ordre, et retourne le premier trouvé — None si absent des
    trois. Utile là où on a juste besoin de l'objet Source/SqlSource/
    WebSource sans savoir a priori son type (routes d'administration
    ciblant une source par son nom).

    ⚠️ À ne pas employer pour décider d'un ACCÈS : un nom absent des
    registres y rend None, ce qui se lit trop facilement comme « aucune
    restriction » alors que ça veut dire « source inconnue » (voir
    _check_doc_access, qui est passé à _searchable_source_names pour
    cette raison)."""
    for registry in (file_sources_config, sql_sources_config, web_sources_config):
        try:
            return registry.get_source(name)
        except KeyError:
            continue
    return None


def _requested_source_names(
    source_names: str | list[str] | None, username: str,
) -> list[str] | None:
    """
    Sources demandées par la requête, restreintes à celles que CET
    utilisateur peut chercher (_searchable_source_names : ni désactivée,
    ni hors de ses groupes, ni absente des registres).

    Retourne None quand rien n'est demandé — « aucun filtre de source »,
    à ne pas confondre avec la liste VIDE, qui dit « tout ce qui était
    demandé a été écarté » et doit alors filtrer sur rien. Sans cette
    distinction, un permalien nommant une source interdite ÉLARGIRAIT la
    recherche à toutes les autres au lieu de ne rien rendre.

    Un nom écarté l'est SILENCIEUSEMENT, et un nom inexistant est traité
    exactement pareil. C'était l'inverse jusqu'ici : 400 « Source
    inconnue » pour un nom absent des trois registres, 200 avec zéro
    résultat pour une source existante mais interdite ou désactivée. La
    différence entre ces deux réponses disait à n'importe quel
    utilisateur si un nom de source existait — un lien profond bricolé à
    la main (`?source=rh-confidentiel`) suffisait à énumérer les sources
    qu'on lui cache. Le prix payé est la détection de faute de frappe,
    que l'aide au zéro résultat rend de toute façon (« sans le filtre
    source : N résultats »).
    """
    if not source_names:
        return None
    names = source_names if isinstance(source_names, list) else [source_names]
    autorisees = set(_searchable_source_names(username))
    return [name for name in names if name in autorisees]


def _visible_to(s, user_groups: list[str]) -> bool:
    """Une source dont allowed_groups est vide est visible par tout le
    monde (comportement historique, avant l'ajout de cette restriction) ;
    sinon il faut être membre d'au moins un des groupes listés."""
    return not s.allowed_groups or any(g in s.allowed_groups for g in user_groups)


def _searchable_source_names(username: str) -> list[str]:
    """
    Noms de TOUTES les sources actuellement cherchables PAR CET
    UTILISATEUR, tous types confondus (fichier/SQL/web) — combine deux
    restrictions indépendantes :
      - "searchable" (voir set_searchable() dans chaque registre) : une
        source peut continuer d'être indexée normalement (watcher/
        sql-worker/web-worker) tout en étant retirée de la consultation.
      - "allowed_groups" (voir set_allowed_groups()) : restreint la
        visibilité de la source aux membres d'un des groupes AD/LDAP
        listés, vide = aucune restriction. Orthogonal à l'ACL par
        document (build_acl_filter) : ceci masque la source en bloc,
        celle-là filtre les documents individuels d'une source par
        ailleurs visible.
    Cette liste est LA définition de « ce que cet utilisateur peut
    atteindre », et les deux chemins d'accès s'y réfèrent :
      - la recherche, fédérée ou non — une source désactivée ou hors
        groupe reste invisible même si elle est explicitement demandée
        via `source` (voir _requested_source_names) ;
      - l'accès direct par doc_id — /document, /api/preview, mots-clés
        personnalisés (voir _check_doc_access).
    Les deux doivent le rester : une restriction que seule la recherche
    applique n'est pas une restriction, juste un tri par défaut.
    """
    user_groups = get_effective_groups(username)
    names = []
    for name, s in file_sources_config.get_sources().items():
        if s.searchable and _visible_to(s, user_groups):
            names.append(name)
    for name, s in sql_sources_config.get_sources().items():
        if s.searchable and _visible_to(s, user_groups):
            names.append(name)
    for name, s in web_sources_config.get_sources().items():
        if s.searchable and _visible_to(s, user_groups):
            names.append(name)
    return names


def _active_custom_facets(source_names: list[str], username: str | None = None) -> dict[str, str]:
    """
    Facettes personnalisées actives pour la recherche en cours —
    {es_field: label} — dérivées des colonnes marquées "facet" dans le
    mapping de CHAQUE source SQL en jeu (voir sql_sources_config.py:
    FieldMapping.facet). Seules les sources SQL peuvent en déclarer :
    fichier/web ont un schéma de document fixe, sans mapping de colonnes.

    "En jeu" = les sources explicitement demandées (`source_names`), ou
    TOUTES les sources SQL cherchables si la recherche est fédérée (pas
    de filtre de source actif) — comme le reste de la recherche
    (_searchable_source_names()). Dédupliqué par nom de champ ES ; en cas
    de collision entre deux sources sur le même nom de champ, le dernier
    libellé rencontré l'emporte.

    `username` est optionnel : None préserve le comportement public de
    GET /custom-facets (aucune restriction par groupe, cohérent avec son
    absence d'authentification) ; les appels depuis /search et
    /search/export le passent pour ne jamais exposer le nom d'une
    facette d'une source hors des groupes autorisés de l'utilisateur.

    ⚠️ Ce `username` ne filtre QUE la liste de repli ci-dessous. Une
    liste `source_names` explicite est prise telle quelle : c'est à
    l'appelant de l'avoir déjà restreinte — ce que fait
    _requested_source_names(), seul chemin par lequel /search et
    /search/export la construisent.
    """
    user_groups = get_effective_groups(username) if username else None
    fallback_sources = sql_sources_config.get_sources().items()
    names = source_names or [
        name for name, s in fallback_sources
        if s.searchable and (user_groups is None or _visible_to(s, user_groups))
    ]
    result: dict[str, str] = {}
    for name in names:
        try:
            source = sql_sources_config.get_source(name)
        except KeyError:
            continue   # source fichier/web, ou nom déjà écarté par _requested_source_names
        for f in source.fields:
            if f.facet:
                result[f.es_field] = f.facet_label or f.es_field
    return result


def _suggestable_custom_facets(username: str) -> dict[str, str]:
    """
    Facettes personnalisées dont les VALEURS peuvent être proposées en
    autocomplétion (voir GET /search/suggest) — {es_field: label}, un
    sous-ensemble de _active_custom_facets().

    Trois retraits, et chacun a sa raison :

    1. **Ce qui n'est pas `keyword`.** La validation admet aussi
       `boolean` en facette (sql_sources_config.py) : agrégeable, certes,
       mais l'`include` d'une agrégation `terms` est une expression
       régulière, qu'Elasticsearch REFUSE sur un champ non textuel. Et
       proposer « true » sous la barre de recherche n'aurait de toute
       façon aucun sens.
    2. **`author` et `keywords`.** Une source SQL a le droit de mapper une
       colonne sur ces noms ; ils sont déjà proposés par le volet fixe
       (user_history.CHAMPS_CORPUS), et les agréger une seconde fois
       afficherait chaque auteur deux fois, sous deux libellés.
    3. Rien d'autre — le contrôle du TYPE RÉELLEMENT EN PLACE dans les
       index appartient à user_history.champs_agregables(), qui interroge
       le moteur. Le type lu ici est celui que la configuration DÉCLARE,
       et les deux divergent dès qu'un index survit à une reconfiguration.
       Ce filtre-ci évite d'interroger le moteur pour un champ dont on
       sait déjà qu'il n'a rien à faire là ; il ne prétend pas trancher.

    Le filtrage par groupes vient de _active_custom_facets(), à qui le
    `username` est passé : le seul NOM d'une facette décrit le schéma
    d'une source, qu'on cache par ailleurs à qui n'y a pas droit. Les
    VALEURS, elles, sont protégées par les filtres passés à
    corpus_terms() — ACL et sources cherchables, ceux de /search.
    """
    # Les types de TOUTES les sources SQL, pas seulement de celles que cet
    # utilisateur voit : deux sources qui déclarent le même nom de champ
    # sous deux types se disputent le mapping de l'alias, et le conflit ne
    # dépend pas de qui regarde. D'où l'ensemble, et l'exigence qu'il ne
    # contienne que `keyword`.
    types: dict[str, set[str]] = {}
    for source in sql_sources_config.get_sources().values():
        for f in source.fields:
            types.setdefault(f.es_field, set()).add(f.es_type)
    fixes = {champ for champ, _ in user_history.CHAMPS_CORPUS}
    return {
        es_field: label
        for es_field, label in _active_custom_facets([], username).items()
        if types.get(es_field) == {"keyword"} and es_field not in fixes
    }


def _collectable_source_names() -> set[str]:
    """
    Noms de TOUTES les sources dont les documents peuvent actuellement
    être ajoutés à une collection ("Mes collections"), tous types
    confondus — indépendant de "searchable" : une source peut rester
    cherchable normalement tout en étant exclue des collections (voir
    set_collectable() dans chaque registre). Utilisé par
    add_collection_document() ; un set (pas une liste) car seul le test
    d'appartenance importe ici, pas l'ordre.
    """
    names = set()
    for name, s in file_sources_config.get_sources().items():
        if s.collectable:
            names.add(name)
    for name, s in sql_sources_config.get_sources().items():
        if s.collectable:
            names.add(name)
    for name, s in web_sources_config.get_sources().items():
        if s.collectable:
            names.add(name)
    return names


@app.get("/searchable-sources")
def get_searchable_sources(user: str = Depends(current_user)):
    """
    Pas d'admin requis (n'importe quel utilisateur identifié peut lister
    les sources qui LUI sont ouvertes) — liste des sources actuellement
    cherchables, pour la présélection de sources AVANT de lancer une
    recherche (voir index.html) — complète la facette "Source" existante,
    qui n'apparaît qu'APRÈS une recherche (dérivée des résultats, avec
    leur compte). Contrairement à /admin/all-sources, pas de nombre de
    documents ni de taille d'index (réservé à l'admin) : juste de quoi
    peupler une liste de cases à cocher. `collectable` est inclus pour la
    même raison que `label`/`type` : index.html s'en sert pour masquer la
    case "ajouter à une collection" sur les résultats d'une source qui
    l'interdit (voir sourceCollectable() côté UI), sans appel séparé.

    Filtrée par allowed_groups (voir _searchable_source_names()) : une
    source restreinte à un groupe dont l'utilisateur n'est pas membre
    n'apparaît pas ici, cohérent avec le fait qu'elle ne renverra de
    toute façon jamais de résultat pour lui dans /search.
    """
    username = user
    user_groups = get_effective_groups(username)
    result = []
    for name, s in file_sources_config.get_sources().items():
        if s.searchable and _visible_to(s, user_groups):
            result.append({"name": name, "label": s.label or name, "type": "file", "collectable": s.collectable})
    for name, s in sql_sources_config.get_sources().items():
        if s.searchable and _visible_to(s, user_groups):
            # card_fields : {champ ES: libellé} pour la carte de résultat.
            # Valeur None = « libellé à dériver du nom », "" = « masquer »,
            # texte = ce libellé. À défaut de card_label, on retombe sur
            # facet_label, déjà saisi pour les colonnes en facette : sans
            # cela, une source configurée avant cette évolution perdrait
            # ses beaux libellés dans la carte.
            card_fields = {
                f.es_field: (f.card_label if f.card_label is not None else f.facet_label)
                for f in s.fields
            }
            result.append({
                "name": name, "label": s.label or name, "type": "sql",
                "collectable": s.collectable, "card_fields": card_fields,
            })
    for name, s in web_sources_config.get_sources().items():
        if s.searchable and _visible_to(s, user_groups):
            result.append({"name": name, "label": s.label or name, "type": "web", "collectable": s.collectable})
    return sorted(result, key=lambda s: s["label"].lower())


@app.get("/custom-facets")
def get_custom_facets(user: str = Depends(current_user)):
    """
    Exige une session (comme tout le reste), mais aucun droit particulier — {es_field: label} de TOUTES les facettes SQL
    personnalisées actives, toutes sources cherchables confondues
    (`_active_custom_facets([])`, même fonction que `/search`, sans
    filtre de source). Contrairement aux facettes renvoyées par
    `/search` (dérivées des résultats d'UNE recherche déjà lancée), ceci
    est disponible dès le chargement de la page — nécessaire pour que la
    syntaxe avancée de la barre de recherche (ex: `bureau:Paris`, voir
    index.html) reconnaisse un opérateur personnalisé avant même la
    première recherche.
    """
    return _active_custom_facets([])


def _resolve_doc_index(doc_id: str) -> str:
    """
    Un doc_id seul ne dit pas dans quel index il vit (recherche
    fédérée) — on le retrouve via une requête `ids` sur l'alias, dont le
    hit renvoie `_index`. Lève 404 si absent de toutes les sources.
    """
    res = es.search(index=ES_SEARCH_ALIAS, query={"ids": {"values": [doc_id]}}, size=1)
    hits = res["hits"]["hits"]
    if not hits:
        raise HTTPException(status_code=404, detail="Document introuvable")
    return hits[0]["_index"]


def _resolve_doc_source(doc_id: str) -> str | None:
    """
    Retourne le nom de la source (champ "source") d'un document, ou None
    s'il est introuvable — utilisé pour vérifier "collectable" avant
    l'ajout à une collection (voir add_collection_document). Ne lève
    jamais d'erreur ici : un doc_id déjà invalide/inaccessible est un cas
    déjà toléré ailleurs par saved_collections.py (une liste peut
    contenir des doc_ids devenus obsolètes).
    """
    try:
        res = es.search(index=ES_SEARCH_ALIAS, query={"ids": {"values": [doc_id]}}, size=1, source=["source"])
        hits = res["hits"]["hits"]
        return hits[0]["_source"].get("source") if hits else None
    except Exception:
        return None


def _journaliser_temps(
    *,
    query: str,
    search_in: str,
    total: int,
    username: str,
    took_ms: int | None,
    duration_ms: float,
    exact: bool = False,
) -> None:
    """Trace le temps d'une recherche dans le journal du service.

    Fonction séparée du endpoint pour rester vérifiable sans monter une
    requête HTTP complète : ce qui se joue ici n'est pas la mesure mais
    la DÉCISION d'écrire ou non, et c'est elle qui doit être testée.

    Au-delà de SLOW_SEARCH_MS, une ligne WARNING avec de quoi rejouer le
    cas (requête, champ, mode, volume, utilisateur). En deçà, une ligne
    DEBUG, muette en exploitation normale, qui permet d'observer une
    période précise en abaissant LOG_LEVEL sans redéployer quoi que ce
    soit.

    `exact` figure dans la trace parce qu'il change les champs
    interrogés (`.exact` au lieu des champs ordinaires) : sans lui, deux
    lignes rigoureusement identiques pourraient décrire deux requêtes ES
    différentes, et « pourquoi celle-ci est-elle lente » n'aurait pas de
    réponse dans le journal.
    """
    if SLOW_SEARCH_MS and duration_ms >= SLOW_SEARCH_MS:
        logger.warning(
            f"[search] Recherche lente : {duration_ms} ms (moteur {took_ms} ms) "
            f"pour '{query}' (search_in={search_in}, exact={exact}, "
            f"{total} résultats, utilisateur {username})"
        )
    else:
        logger.debug(
            f"[search] {duration_ms} ms (moteur {took_ms} ms) "
            f"pour '{query}' ({total} résultats)"
        )


# ── Recherche ────────────────────────────────────────────────
# ── Aide au zéro résultat ───────────────────────────────────────
#
# L'administration disposait déjà de la liste des recherches
# infructueuses (voir /admin/search-logs/zero-results) ; l'utilisateur,
# lui, avait un écran vide et une phrase. Le diagnostic existait, l'aide
# non.
#
# Trois pistes, dans l'ordre où elles servent : la correction
# orthographique, le relâchement d'un filtre (« 12 résultats sans le
# filtre .pdf »), et les autres sources.
#
# ⚠️ Deux règles gouvernent tout ce qui suit :
#
# 1. **Chaque compte annoncé doit être atteignable.** Les filtres ACL et
#    de sources cherchables ne sont JAMAIS relâchés : annoncer « 12
#    résultats » puis afficher une liste vide après le clic coûte plus de
#    confiance qu'un écran vide honnête.
# 2. **Rien de tout ceci ne s'exécute sur le chemin nominal.** Une
#    recherche qui trouve quelque chose ne paie pas un millième de
#    seconde pour ce code.

# Au-delà, l'écran devient une liste de courses plutôt qu'une aide.
MAX_RELACHEMENTS = 6


def _corriger_requete(texte: str, entrees: list) -> str | None:
    """Rebâtit la requête à partir des corrections proposées par
    Elasticsearch, en respectant les positions qu'il rapporte — un
    remplacement naïf par `str.replace` toucherait aussi les occurrences
    d'un mot dans un autre."""
    morceaux, curseur = [], 0
    for entree in entrees:
        options = entree.get("options") or []
        if not options:
            continue
        debut, longueur = entree["offset"], entree["length"]
        morceaux.append(texte[curseur:debut])
        morceaux.append(options[0]["text"])
        curseur = debut + longueur
    if not morceaux:
        return None
    morceaux.append(texte[curseur:])
    corrigee = "".join(morceaux).strip()
    return corrigee if corrigee.casefold() != texte.casefold() else None


def _aide_zero_resultat(
    *, must: list, obligatoires: list, relachables: dict,
    facet_filters: dict, fields: list, query_text: str,
) -> dict:
    """Ce qu'on peut proposer à quelqu'un dont la recherche n'a rien
    donné. Meilleur effort : toute panne d'Elasticsearch ici rend un
    objet vide, l'écran retombant sur le message d'origine."""
    droppables = {nom: f for nom, f in {**facet_filters, **relachables}.items() if f}

    recherches: list[dict] = []
    plans: list[str] = []

    def ajouter(plan: str, corps: dict) -> None:
        plans.append(plan)
        recherches.append({})
        recherches.append(corps)

    def compter(filtres: list) -> dict:
        return {"size": 0, "query": {"bool": {"must": must, "filter": filtres}}}

    for nom in list(droppables)[:MAX_RELACHEMENTS]:
        autres = [f for autre, f in droppables.items() if autre != nom]
        ajouter(f"sans:{nom}", compter(obligatoires + autres))

    # « Sans aucun filtre » n'a de sens qu'à partir de deux : avec un
    # seul, cette ligne dirait exactement la même chose que la précédente.
    if len(droppables) > 1:
        ajouter("sans:__all__", compter(obligatoires))

    if "source" in droppables:
        autres = [f for autre, f in droppables.items() if autre != "source"]
        corps = compter(obligatoires + autres)
        corps["aggs"] = {"sources": {"terms": {"field": "source", "size": 10}}}
        ajouter("sources", corps)

    if query_text:
        ajouter("suggestion", {
            "size": 0,
            "suggest": {
                "texte": {
                    "text": query_text,
                    # "popular" : ne propose qu'un terme PLUS fréquent que
                    # celui tapé. Sans ça, le correcteur propose volontiers
                    # une variante aussi rare que l'original, ce qui ne
                    # corrige rien.
                    "term": {"field": "content", "suggest_mode": "popular", "size": 1},
                }
            },
        })

    if not recherches:
        return {}

    try:
        reponses = es.msearch(index=ES_SEARCH_ALIAS, searches=recherches)["responses"]
    except Exception as e:
        logger.warning(f"[zero-result] Aide indisponible pour « {query_text} » : {e}")
        return {}

    aide: dict = {"relaxations": [], "sources": [], "suggestion": None}
    # strict=False : Elasticsearch renvoie une réponse par recherche
    # envoyée, mais une aide au zéro résultat n'a pas à faire échouer la
    # recherche elle-même si cette invariante venait à bouger.
    for plan, reponse in zip(plans, reponses, strict=False):
        if reponse.get("error"):
            continue
        if plan.startswith("sans:"):
            total = reponse["hits"]["total"]["value"]
            if total:
                aide["relaxations"].append({"field": plan[5:], "count": total})
        elif plan == "sources":
            aide["sources"] = reponse.get("aggregations", {}).get("sources", {}).get("buckets", [])
        elif plan == "suggestion":
            entrees = reponse.get("suggest", {}).get("texte", [])
            aide["suggestion"] = _corriger_requete(query_text, entrees)

    # ⚠️ Le correcteur travaille sur le dictionnaire de termes de l'index,
    # que l'ACL ne filtre pas : un mot tiré d'un document interdit
    # pourrait être proposé. On ne rend donc la correction que si elle
    # donne des résultats VISIBLES PAR CET UTILISATEUR — ce qui la
    # débarrasse du même coup des corrections qui ne mènent nulle part.
    if aide["suggestion"]:
        try:
            visibles = es.count(
                index=ES_SEARCH_ALIAS,
                query={
                    "bool": {
                        "must": [{"multi_match": {
                            "query": aide["suggestion"], "fields": fields, "fuzziness": "AUTO",
                        }}],
                        "filter": obligatoires,
                    }
                },
            )["count"]
        except Exception:
            visibles = 0
        if not visibles:
            aide["suggestion"] = None

    return aide if (aide["suggestion"] or aide["relaxations"] or aide["sources"]) else {}


def _documents_epingles(requete: str, username: str) -> list[dict]:
    """Les documents épinglés sur cette requête, tels que CET utilisateur
    a le droit de les voir.

    ⚠️ Relus par une vraie recherche portant le filtre ACL et la
    restriction aux sources cherchables — jamais par un `mget`, qui
    rendrait le document sans rien vérifier. Un épinglage met en avant,
    il n'autorise pas : celui qui n'a pas le droit de voir le document ne
    le voit pas, et rien à l'écran ne lui apprend qu'il existe.

    Un document épinglé puis supprimé de l'index disparaît de lui-même :
    la recherche ne le trouve plus, la liste se réduit. Le panneau
    d'administration, lui, le signale comme introuvable pour qu'on puisse
    nettoyer la règle.

    Meilleur effort : toute erreur ici rend une liste vide plutôt que de
    faire échouer la recherche.
    """
    identifiants = pinned.pour_requete(requete)
    if not identifiants:
        return []
    try:
        res = es.search(
            index=ES_SEARCH_ALIAS,
            size=len(identifiants),
            query={"bool": {"filter": [
                {"ids": {"values": identifiants}},
                build_acl_filter(username),
                {"terms": {"source": _searchable_source_names(username)}},
            ]}},
            source_excludes=["content", "content_vector"],
        )
    except Exception as e:
        logger.warning(f"[epingles] Relecture impossible pour « {requete} » : {e}")
        return []

    trouves = {h["_id"]: h["_source"] for h in res["hits"]["hits"]}
    # L'ordre est celui de l'administration, pas celui d'Elasticsearch :
    # quand quelqu'un épingle trois documents, il les a classés.
    return [
        {**trouves[identifiant], "id": identifiant, "score": None, "highlight": [], "pinned": True}
        for identifiant in identifiants
        if identifiant in trouves
    ]


def _config_surlignage(exact: bool) -> tuple[str, dict]:
    """Champ à surligner pour l'extrait, et les options du surlignage.

    Source unique de /search et de /search/export, qui portaient chacun
    leur copie de ces réglages.

    LE CHAMP SUIT LA CLAUSE. En recherche exacte, celle-ci vise
    `content.exact`, analysé sans racinisation ; surligner `content`,
    analysé en français, revient à chercher les jetons de l'un dans ceux
    de l'autre. Les termes extraits de la requête sont les formes exactes
    (« annuelles »), les jetons du champ sont racinisés (« annuel ») :
    aucune intersection, donc ES ne trouve AUCUN passage et ne renvoie
    PAS DE FRAGMENT — pas un extrait sans marques, pas d'extrait du tout.
    La carte de résultat n'affichait alors rien sous le titre.

    Le défaut ne se voyait pas sur tous les mots, d'où son aspect
    aléatoire : il disparaît dès que le radical est égal à la forme
    exacte (« budget » → « budget »), et c'est sur ce genre de requête
    qu'il avait été cru corrigé. Mesuré sur le corpus de développement
    avant correction : « budget » rendait ses extraits, « délégations »,
    « millions », « annuelles » et « systèmes » n'en rendaient aucun.

    `require_field_match` reste à son défaut (True), et le `False` posé
    ici auparavant disparaît : il n'était que la tentative de faire
    surligner `content` par une clause portant sur `content.exact`, ce
    qui ne pouvait pas fonctionner (voir ci-dessus). Le défaut est de
    surcroît le comportement voulu — un terme trouvé via le titre ou
    l'auteur n'a jamais eu à être surligné dans le corps.

    Un index qui n'a pas reçu la migration (`./manage.sh migrer-exact
    --apply`) n'a pas de sous-champ `content.exact` : ES ignore
    silencieusement un champ de surlignage absent du mapping, sans
    erreur ni shard en échec. Ces index-là restent donc muets en extrait
    exactement comme ils le sont déjà en recherche exacte — le
    `multi_match` n'y matche rien non plus.
    """
    return (
        "content.exact" if exact else "content",
        {
            "fragment_size":        200,
            "number_of_fragments":  2,
            # max_analyzed_offset : sans lui, dès qu'un document du lot
            # (ex: gros PST/PDF) dépasse index.highlight.max_analyzed_offset
            # (1 000 000 caractères), ES fait échouer TOUS les shards
            # portant ce document, et le highlighting renvoie alors
            # hits.hits=[] pour la requête entière (total correct, mais
            # aucun résultat) — d'où des recherches qui semblaient soudain
            # ne plus rien retourner. On tronque explicitement l'analyse du
            # surlignage à cette limite.
            "max_analyzed_offset":  1000000,
        },
    )


def _extraits(hit: dict) -> list[str]:
    """Fragments surlignés d'un hit, quel que soit le champ dont ils
    viennent.

    ES range les fragments sous le NOM DU CHAMP surligné : `content` en
    recherche ordinaire, `content.exact` en recherche exacte. Les lire
    sous une clé écrite en dur rendrait la liste vide dans l'un des deux
    modes — le défaut que corrige _config_surlignage(), simplement déplacé
    d'un cran.

    D'où la lecture des deux clés plutôt que le passage du champ retenu
    jusqu'ici : une seule des deux peut être présente (une requête ne
    surligne qu'un champ), et aucun appelant ne peut se désynchroniser de
    _config_surlignage() en oubliant de propager son choix.

    Le format rendu, lui, ne dépend pas du mode : une liste plate de
    fragments, sous « highlight » pour l'interface et concaténée dans la
    colonne « Extrait » de l'export.
    """
    surlignage = hit.get("highlight", {})
    return surlignage.get("content") or surlignage.get("content.exact") or []


def _verifier_shards(res: dict, contexte: str) -> None:
    """Refuse une réponse Elasticsearch PARTIELLE plutôt que de la lire
    comme un résultat.

    ES n'échoue pas quand une partie seulement des shards interrogés
    échoue : `allow_partial_search_results` vaut vrai par défaut, il
    répond 200 avec les shards survivants et range les autres sous
    `_shards.failures`. `hits.total` et les agrégations ne portent alors
    que sur les survivants — indiscernable, pour l'appelant, d'un corpus
    qui ne contient rien.

    Cas vécu (2026-08-13) : une facette personnalisée déclarée sur un
    champ `text` (voir _active_custom_facets) posait une agrégation
    `terms` sur un champ sans doc_values. 13 shards sur 14 échouaient, et
    la recherche fédérée répondait « 0 résultat » sur un corpus de 23 000
    documents — en 5 ms, sans erreur, sans une ligne de log, toutes
    facettes vides. Une recherche restreinte à une source sans facette
    personnalisée répondant normalement, le défaut a été cherché du côté
    des ACL, de l'annuaire et du registre des sources avant d'être trouvé
    là.

    D'où le refus plutôt que le rendu partiel assorti d'un avertissement :
    un compte faux coûte plus cher qu'une erreur, précisément parce qu'il
    ne se remarque pas. `contexte` nomme la route pour que le journal
    distingue /search de /search/export.
    """
    shards = res.get("_shards") or {}
    echecs = shards.get("failed") or 0
    if not echecs:
        return

    # Une même cause frappe en général tous les shards touchés (un champ
    # non agrégeable l'est dans chaque index) : on remonte le premier
    # motif et le nombre de shards, pas N fois le même message.
    failures = shards.get("failures") or []
    premier  = (failures[0] if failures else {}).get("reason") or {}
    motif    = premier.get("reason") or premier.get("type") or "cause non précisée par Elasticsearch"
    index    = (failures[0] if failures else {}).get("index") or "?"

    detail = (
        f"Résultat incomplet : {echecs} shard(s) sur {shards.get('total')} en échec "
        f"— les documents qu'ils portent sont absents du compte et des facettes. "
        f"Premier motif ([{index}]) : {motif}"
    )
    logger.error(f"[{contexte}] {detail}")
    raise HTTPException(status_code=500, detail=detail)


@app.post("/search")
def search(
    req: SearchQuery,
    request: Request,
    user: str = Depends(current_user),
):
    # Départ du chronomètre AVANT toute autre chose : la résolution des
    # groupes ACL et la relecture de la configuration des champs sont
    # elles aussi du temps que l'utilisateur attend. perf_counter et non
    # time.time() — seul le premier est monotone, donc insensible à un
    # ajustement d'horloge (NTP) pendant la mesure.
    t0 = time.perf_counter()

    _ensure_index_exists()
    username   = user
    acl_filter = build_acl_filter(username)

    # Champs, poids et clause de texte : source unique dans
    # search_query.py, partagée avec l'export et la vérification
    # d'alertes. Les poids étant réglables depuis l'administration, une
    # copie littérale ici divergerait du classement affiché dès le
    # premier réglage — et la même remarque vaut désormais pour les
    # règles de guillemets/flou/exactitude.
    #
    # Une requête vide (champ de recherche vide mais filtres actifs — ex:
    # syntaxe avancée "auteur:...", "type:..." utilisée seule) y donne un
    # match_all : les filtres ci-dessous restent alors seuls responsables
    # de la restriction.
    query_text = req.query.strip()
    is_exact_phrase = search_query.est_phrase(query_text)
    must = [search_query.build_text_clause(query_text, req.search_in, req.exact)]

    # Même jeu de champs que la clause ci-dessus (d'où le même `exact`) :
    # l'aide au zéro résultat compte des résultats qu'elle propose ensuite
    # d'aller voir, et annoncer « 12 résultats sans ce filtre » sur des
    # champs que la recherche n'interroge pas mènerait à un écran vide.
    sets = search_query.field_sets(exact=req.exact)
    fields = sets.get(req.search_in, sets["all"])

    # Filtres "de base" : toujours appliqués, jamais concernés par
    # l'exclusion décrite ci-dessous (ACL, pièces jointes, période,
    # sources désactivées pour la recherche). Une source "searchable:
    # false" (voir set_searchable()) est retirée ICI, en amont de tout —
    # donc invisible même si explicitement demandée via `source` : la
    # désactivation est absolue, pas seulement "absente par défaut".
    # Les deux premiers ne se relâchent JAMAIS : ils disent ce que cet
    # utilisateur a le droit de voir, et l'aide au zéro résultat (plus
    # bas) doit rester dedans — proposer « 12 résultats sans ce filtre »
    # pour des documents qu'il ne pourra pas ouvrir coûte plus de
    # confiance qu'un écran vide honnête.
    filtres_obligatoires = [
        acl_filter,   # ACL en premier — mis en cache par ES
        {"terms": {"source": _searchable_source_names(username)}},
    ]
    # Ceux-là s'appliquent aussi à toute recherche, mais viennent d'un
    # choix de l'utilisateur : ils peuvent donc être proposés au
    # relâchement quand la recherche ne donne rien.
    filtres_relachables = {}
    if req.has_attachments:
        filtres_relachables["has_attachments"] = {"term": {"has_attachments": True}}
    if req.date_from or req.date_to:
        r = {}
        if req.date_from: r["gte"] = req.date_from
        if req.date_to:   r["lte"] = req.date_to
        filtres_relachables["date"] = {"range": {"date_modified": r}}

    base_filters = filtres_obligatoires + list(filtres_relachables.values())

    # Filtres "de facette" : chacun correspond à une agrégation affichée
    # dans la barre latérale (extension/auteur/dossier/source), à
    # sélection cumulative en OU — sauf les mots-clés, seule facette
    # multi-valuée, combinée en ET (voir _keywords_filter).
    # Construits à part des base_filters pour
    # pouvoir, plus bas, calculer le compte de chaque facette en
    # EXCLUANT le filtre de cette facette elle-même — sinon, sélectionner
    # un premier auteur ferait disparaître tous les autres de la liste
    # (impossible d'en cocher un second), pareil pour source/dossier.
    # Motif standard de "faceted navigation" avec post_filter + filter
    # aggregations : https://www.elastic.co/guide/en/elasticsearch/reference/current/search-aggregations-bucket-terms-aggregation.html
    extension_filter = None
    if req.extension:
        # Valeur(s) brutes du champ ES, point compris (".pdf", ".docx"...)
        # — même format que les clés retournées par facets.extensions,
        # pas de transformation ici (même principe que author/source :
        # le client envoie exactement ce que la facette lui a donné).
        exts = req.extension if isinstance(req.extension, list) else [req.extension]
        extension_filter = {"terms": {"extension": exts}}

    author_filter = None
    if req.author:
        authors = req.author if isinstance(req.author, list) else [req.author]
        author_filter = {"terms": {"author": authors}}

    # Seule facette combinée en ET (intersection) — voir _keywords_filter.
    keywords_filter = _keywords_filter(req.keywords)

    folder_filter = _folder_filter(req.folder)

    # None = aucune source demandée ; liste vide = tout ce qui était
    # demandé a été écarté, et le filtre doit alors ne rien matcher (voir
    # _requested_source_names) — d'où le test sur None et non sur la
    # vacuité.
    source_names  = _requested_source_names(req.source, username)
    source_filter = {"terms": {"source": source_names}} if source_names is not None else None

    # Facettes personnalisées (voir _active_custom_facets) : une par champ
    # marqué "facet" sur l'une des sources SQL actuellement en jeu. Les
    # clés de req.custom qui ne correspondent à aucune facette active sont
    # silencieusement ignorées (ex: source changée depuis, valeur
    # obsolète restée dans le state du navigateur) plutôt que de lever une
    # erreur — même tolérance que le reste de la recherche vis-à-vis d'un
    # état client périmé.
    custom_facet_defs = _active_custom_facets(source_names or [], username)
    custom_filters = {}
    for es_field in custom_facet_defs:
        values = (req.custom or {}).get(es_field)
        if values:
            custom_filters[f"custom:{es_field}"] = {"terms": {es_field: values}}

    facet_filters = {
        "extension": extension_filter,
        "author":    author_filter,
        "keywords":  keywords_filter,
        "folder":    folder_filter,
        "source":    source_filter,
        **custom_filters,
    }
    active_facet_filters = [f for f in facet_filters.values() if f]

    def facet_agg(field: str, size: int, exclude: str | None) -> dict:
        """Agrégation de facette qui exclut son propre filtre (voir plus
        haut) mais applique tous les AUTRES filtres de facette actifs —
        cocher un auteur ne doit réduire que les dossiers/sources/
        extensions affichés, jamais la liste des autres auteurs.

        exclude=None n'exclut rien, donc TOUS les filtres actifs
        s'appliquent : réservé aux mots-clés, combinés en ET. Là, le
        compte affiché à côté d'une valeur est exactement le nombre de
        résultats obtenu en la cochant, et les valeurs qui mèneraient à
        zéro disparaissent au lieu d'être proposées en vain."""
        others = [f for name, f in facet_filters.items() if f and name != exclude]
        return {
            "filter": {"bool": {"filter": others}} if others else {"match_all": {}},
            "aggs":   {"values": {"terms": {
                "field": field,
                "size":  size,
                # Sans shard_size explicite, ES ne remonte que ses
                # ~40 premiers termes PAR SHARD (22 ici) et additionne :
                # une valeur absente du haut de tableau d'un shard voit
                # sa part perdue, donc un compte SOUS-ESTIMÉ (mesuré :
                # 30 affichés pour 37 documents réels). Inacceptable
                # pour les mots-clés, dont le compte annonce désormais
                # le nombre de résultats après clic — et trompeur
                # ailleurs. 500 ramène doc_count_error_upper_bound à 0
                # sur le corpus actuel, pour un coût mémoire négligeable
                # (22 × 500 termes agrégés).
                "shard_size": FACET_SHARD_SIZE,
            }}},
        }

    sort_clause = (
        [{"_score": "desc"}]
        if req.sort == "_score"
        # "missing": "_last" explicite plutôt que de compter sur le
        # comportement par défaut d'ES — utile ici car les emails PST
        # n'ont pas de champ "size" (pst_extractor.py ne l'indexe pas),
        # donc un tri par taille doit gérer ces valeurs absentes.
        else [{req.sort: {"order": "desc", "missing": "_last"}}, {"_score": "desc"}]
    )

    # Le champ surligné suit les champs interrogés — voir _config_surlignage.
    champ_extrait, options_extrait = _config_surlignage(req.exact)

    try:
        res = es.search(
            index=ES_SEARCH_ALIAS,
            query={"bool": {"must": must, "filter": base_filters}},
            # Les filtres de facette s'appliquent aux résultats ICI (via
            # post_filter, évalué après les agrégations) plutôt que dans
            # `query` — c'est ce qui permet à chaque facet_agg() ci-dessus
            # de les exclure sélectivement sans que les résultats
            # eux-mêmes cessent de respecter TOUS les filtres actifs.
            post_filter={"bool": {"filter": active_facet_filters}},
            highlight={
                "fields": {champ_extrait: options_extrait},
                # Sans ceci, ES utilise ses balises par défaut (<em>...</em>),
                # qui ne correspondent à AUCUNE règle CSS du frontend — les
                # termes trouvés n'étaient donc jamais visuellement surlignés,
                # juste en italique. On lui fait directement émettre la classe
                # CSS attendue.
                "pre_tags":  ['<mark class="highlight">'],
                "post_tags": ["</mark>"],
            },
            sort=sort_clause,
            # Nécessaire pour que le tri secondaire par _score (utilisé
            # comme départage quand le tri principal n'est pas _score)
            # soit réellement calculé — sans ça, ES ne calcule pas les
            # scores du tout en dehors d'un tri _score primaire.
            track_scores=True,
            from_=req.from_,
            size=req.size,
            # Exclusion plutôt qu'une liste fixe de champs à inclure : une
            # liste fixe (motif précédent) cassait silencieusement tout
            # champ personnalisé mappé par une source SQL/web (ex: colonne
            # "telephone" de la source "agents", voir sql_sources_config.py)
            # — jamais renvoyé au navigateur, sans erreur visible.
            # "content"/"content_vector" restent exclus (volumineux, jamais
            # affichés directement dans un résultat de recherche).
            source_excludes=["content", "content_vector"],
            aggs={
                "by_extension": facet_agg("extension",  10, "extension"),
                "by_author":    facet_agg("author",     10, "author"),
                # Seule facette sans auto-exclusion : en ET, les valeurs
                # utiles sont celles qui CO-OCCURRENT avec la sélection.
                "by_keywords":  facet_agg("keywords",   20, None),
                "by_folder":    facet_agg("folder_top", 10, "folder"),
                "by_source":    facet_agg("source",     20, "source"),
                **{
                    f"by_custom__{es_field}": facet_agg(es_field, 20, f"custom:{es_field}")
                    for es_field in custom_facet_defs
                },
            }
        )
    except Exception as e:
        # Remonte le vrai message ES plutôt qu'un 500 générique opaque
        # ("Internal Server Error") — indispensable pour diagnostiquer
        # un problème de tri/requête sans avoir à fouiller les logs.
        # La durée écoulée distingue une requête refusée d'emblée (erreur
        # de syntaxe, quelques millisecondes) d'un timeout au bout des 60
        # secondes de request_timeout — deux pannes sans rapport.
        logger.error(
            f"[search] Erreur ES pour la requête '{req.query}' (sort={req.sort}) "
            f"après {round((time.perf_counter() - t0) * 1000)} ms : {e}"
        )
        raise HTTPException(status_code=400, detail=f"Erreur de recherche : {e}") from e

    # AVANT toute lecture de la réponse : un shard en échec fait mentir à
    # la fois `total`, les résultats de la page et toutes les facettes.
    _verifier_shards(res, "search")

    hits  = res["hits"]["hits"]
    total = res["hits"]["total"]["value"]

    # Temps passé DANS Elasticsearch, tel qu'il le rapporte lui-même :
    # n'inclut ni le réseau, ni la sérialisation, ni tout ce que fait
    # l'API autour. C'est ce qui permet de trancher entre « le moteur est
    # lent » et « l'API est lente » — la durée totale seule ne le dit pas.
    took_ms = res.get("took")

    # Arrêt du chronomètre AVANT la journalisation : l'écriture du log de
    # recherche dans ES n'est pas du temps de recherche, et l'inclure
    # ferait passer une panne du journal pour une lenteur du moteur.
    # L'assemblage de la réponse ci-dessous reste également hors mesure,
    # mais il ne fait que parcourir une page de résultats déjà en mémoire.
    duration_ms = round((time.perf_counter() - t0) * 1000, 1)

    _journaliser_temps(
        query=req.query,
        search_in=req.search_in,
        exact=req.exact,
        total=total,
        username=username,
        took_ms=took_ms,
        duration_ms=duration_ms,
    )

    search_id = search_log.log_search(
        es,
        username=username,
        # get_user_groups est en @lru_cache : cet appel à chaque recherche
        # ne déclenche pas un aller-retour LDAP à chaque fois.
        groups=get_effective_groups(username),
        ip=get_client_ip(request),
        query=req.query,
        search_in=req.search_in,
        source=req.source,
        total_results=total,
        result_files=[h["_source"].get("filename", "") for h in hits],
        extension=req.extension,
        author=req.author,
        folder=req.folder,
        keywords=req.keywords,
        date_from=req.date_from,
        date_to=req.date_to,
        took_ms=took_ms,
        duration_ms=duration_ms,
    )

    # Uniquement quand la recherche n'a rien donné : le chemin nominal ne
    # paie rien pour cette aide (voir _aide_zero_resultat).
    aide = _aide_zero_resultat(
        must=must,
        obligatoires=filtres_obligatoires,
        relachables=filtres_relachables,
        facet_filters=facet_filters,
        fields=fields,
        query_text=query_text[1:-1].strip() if is_exact_phrase else query_text,
    ) if total == 0 else {}

    # Épinglés : uniquement sur la première page. Les répéter en tête de
    # chaque page ferait passer l'utilisateur devant les mêmes documents
    # à chaque « suivant », en croyant les avoir déjà dépassés.
    epingles = _documents_epingles(req.query, username) if req.from_ == 0 else []
    ids_epingles = {document["id"] for document in epingles}

    return {
        "total":     total,
        "username":  username,
        "search_id": search_id,
        # Absent tant qu'il y a des résultats, et absent aussi quand on
        # n'a rien à proposer : l'interface n'affiche ce bloc que s'il
        # existe.
        **({"zero_result": aide} if aide else {}),
        # Un sous-objet plutôt que deux clés à plat : l'interface n'a
        # qu'une seule chose à tester pour savoir si elle a une mesure à
        # afficher, et les deux durées restent visiblement solidaires.
        "timing": {
            "took_ms":     took_ms,
            "duration_ms": duration_ms,
        },
        "results": [
            {
                **h["_source"],
                # APRÈS le dépaquetage, jamais avant : une source SQL qui
                # projette sa colonne "id" dans un champ ES du même nom
                # écrasait sinon l'identifiant Elasticsearch par la clé
                # primaire SQL. Tout ce qui s'appuie ensuite sur cet
                # identifiant — /document/{id}, les collections, le suivi
                # de clic — visait alors un document inexistant.
                "id":        h["_id"],
                "score":     round(h["_score"], 4),
                "highlight": _extraits(h),
            }
            for h in hits
            # Un document épinglé qui figure aussi dans les résultats
            # naturels n'est pas affiché deux fois : il est rendu une
            # seule fois, dans le bloc épinglé. Le total, lui, ne bouge
            # pas — il compte des documents trouvés, pas des cartes.
            if h["_id"] not in ids_epingles
        ],
        # Absent quand il n'y a rien à épingler : l'interface n'affiche
        # ce bloc que s'il existe.
        **({"pinned": epingles} if epingles else {}),
        "facets": {
            "extensions": res["aggregations"]["by_extension"]["values"]["buckets"],
            "authors":    res["aggregations"]["by_author"]["values"]["buckets"],
            "keywords":   res["aggregations"]["by_keywords"]["values"]["buckets"],
            "folders":    res["aggregations"]["by_folder"]["values"]["buckets"],
            "sources":    res["aggregations"]["by_source"]["values"]["buckets"],
            # Facettes propres à la ou aux source(s) SQL en jeu (ex: "Bureau"/
            # "Fonction" pour la source "agents") — absentes/vides tant
            # qu'aucune source n'en déclare (voir _active_custom_facets()) ;
            # index.html construit/retire dynamiquement leur section de
            # sidebar à partir de cette clé.
            "custom": {
                es_field: {
                    "label":   label,
                    "buckets": res["aggregations"][f"by_custom__{es_field}"]["values"]["buckets"],
                }
                for es_field, label in custom_facet_defs.items()
            },
        }
    }


# ── Export des résultats de recherche (XLSX / DOCX) ─────────────
# Même critères que POST /search (SearchQuery), mais TOUS les résultats
# correspondants (jusqu'à SEARCH_EXPORT_MAX_ROWS) plutôt que la seule
# page affichée — d'où une requête ES séparée, sans les agrégations de
# facettes (inutiles ici, pas d'UI à peupler).
SEARCH_EXPORT_MAX_ROWS = 500


class SearchExportQuery(SearchQuery):
    format: str = "xlsx"   # "xlsx" | "docx"


def _build_search_query(req: SearchQuery, username: str) -> dict:
    """
    Construit la requête ES (must + filtres) pour une recherche —
    factorisé entre POST /search et POST /search/export. Ne couvre PAS
    les agrégations de facettes ni le post_filter associé (spécifiques
    à /search, sans objet pour un export).
    """
    acl_filter = build_acl_filter(username)
    # Champs, poids et règles de guillemets/flou/exactitude : source
    # unique dans search_query.py, partagée avec /search et la
    # vérification d'alertes. Une copie littérale ici divergerait du
    # classement affiché dès le premier réglage des poids depuis
    # l'administration — un export doit contenir ce que l'écran montrait.
    must = [search_query.build_text_clause(req.query.strip(), req.search_in, req.exact)]

    filters = [acl_filter, {"terms": {"source": _searchable_source_names(username)}}]
    if req.has_attachments:
        filters.append({"term": {"has_attachments": True}})
    if req.date_from or req.date_to:
        r = {}
        if req.date_from: r["gte"] = req.date_from
        if req.date_to:   r["lte"] = req.date_to
        filters.append({"range": {"date_modified": r}})
    if req.extension:
        exts = req.extension if isinstance(req.extension, list) else [req.extension]
        filters.append({"terms": {"extension": exts}})
    if req.author:
        authors = req.author if isinstance(req.author, list) else [req.author]
        filters.append({"terms": {"author": authors}})
    keywords_filter = _keywords_filter(req.keywords)   # ET, pas OU — voir le helper
    if keywords_filter:
        filters.append(keywords_filter)
    folder_filter = _folder_filter(req.folder)
    if folder_filter:
        filters.append(folder_filter)
    # Voir le commentaire équivalent dans /search : None et liste vide ne
    # veulent pas dire la même chose.
    source_names = _requested_source_names(req.source, username)
    if source_names is not None:
        filters.append({"terms": {"source": source_names}})
    for es_field, values in (req.custom or {}).items():
        # Même tolérance qu'en recherche normale (voir _active_custom_facets
        # dans /search) : une clé qui ne correspond plus à une facette
        # active de la/des source(s) en jeu est ignorée plutôt que rejetée.
        if values and es_field in _active_custom_facets(source_names or [], username):
            filters.append({"terms": {es_field: values}})

    return {"bool": {"must": must, "filter": filters}}


def _slugify_query(text: str) -> str:
    """Nom de fichier sûr à partir de la requête — alphanumérique et
    tirets seulement, tronqué pour rester raisonnable."""
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", text.strip().lower()).strip("-")
    return (slug or "recherche")[:60]


def _export_results_xlsx(query_text: str, hits: list) -> StreamingResponse:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Résultats de recherche"
    ws.append(["Nom", "Extension", "Auteur", "Mots-clés", "Source", "Dossier",
               "Date de modification", "Taille (o)", "Chemin", "Extrait"])
    for h in hits:
        s = h["_source"]
        snippet = " … ".join(_extraits(h))
        ws.append([
            s.get("filename", ""),
            s.get("extension", ""),
            s.get("author", ""),
            ", ".join(s.get("keywords") or []),
            s.get("source", ""),
            s.get("folder", ""),
            s.get("date_modified", ""),
            s.get("size", 0),
            s.get("filepath", ""),
            snippet,
        ])
    for col_idx, width in enumerate([32, 10, 18, 24, 14, 24, 18, 12, 50, 60], start=1):
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    filename = f"resultats-{_slugify_query(query_text)}.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _export_results_docx(query_text: str, hits: list) -> StreamingResponse:
    from docx import Document

    doc = Document()
    doc.add_heading(f'Résultats de recherche — « {query_text} »', level=1)
    doc.add_paragraph(f"{len(hits)} document(s)")
    for h in hits:
        s = h["_source"]
        doc.add_heading(s.get("filename") or "(sans nom)", level=2)
        meta = []
        if s.get("author"):        meta.append(f"Auteur : {s['author']}")
        if s.get("keywords"):      meta.append(f"Mots-clés : {', '.join(s['keywords'])}")
        if s.get("source"):        meta.append(f"Source : {s['source']}")
        if s.get("folder"):        meta.append(f"Dossier : {s['folder']}")
        if s.get("date_modified"): meta.append(f"Modifié le : {s['date_modified'][:10]}")
        if meta:
            doc.add_paragraph(" · ".join(meta))
        if s.get("filepath"):
            doc.add_paragraph(s["filepath"], style="Intense Quote")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"resultats-{_slugify_query(query_text)}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/search/export")
def export_search_results(req: SearchExportQuery, user: str = Depends(current_user)):
    """
    Export XLSX ou DOCX des résultats d'une recherche — mêmes critères
    que POST /search (même corps de requête, avec juste "format" en
    plus), mais jusqu'à SEARCH_EXPORT_MAX_ROWS résultats plutôt que la
    seule page affichée à l'écran.
    """
    if not ui_config.get_config().get("export_enabled", True):
        raise HTTPException(status_code=403, detail="L'export des résultats est désactivé.")
    _ensure_index_exists()
    username = user
    query = _build_search_query(req, username)

    # Le champ surligné suit les champs interrogés — voir _config_surlignage.
    champ_extrait, options_extrait = _config_surlignage(req.exact)

    sort_clause = (
        [{"_score": "desc"}]
        if req.sort == "_score"
        else [{req.sort: {"order": "desc", "missing": "_last"}}, {"_score": "desc"}]
    )

    try:
        res = es.search(
            index=ES_SEARCH_ALIAS,
            query=query,
            sort=sort_clause,
            track_scores=True,
            size=SEARCH_EXPORT_MAX_ROWS,
            # Voir le commentaire équivalent dans /search — exclusion plutôt
            # qu'une liste figée, pour la même raison.
            source_excludes=["content", "content_vector"],
            highlight={
                # Mêmes champ et options qu'à l'écran (dont
                # max_analyzed_offset, qui évite ici qu'un seul document
                # trop long parmi les 500 lignes fasse échouer le
                # highlighting sur tous les shards qui le portent et
                # rende un export à 0 ligne) : la colonne « Extrait »
                # doit dire ce que la carte de résultat affichait.
                "fields": {champ_extrait: options_extrait},
                # Pas de balises de surlignage ici (texte brut pour un export,
                # contrairement à /search qui les affiche en HTML).
                "pre_tags": [""], "post_tags": [""],
            },
        )
    except Exception as e:
        logger.error(f"[search/export] Erreur ES pour la requête '{req.query}' : {e}")
        raise HTTPException(status_code=400, detail=f"Erreur de recherche : {e}") from e

    # Même règle qu'à l'écran (voir _verifier_shards), et le silence y
    # serait pire : un tableur amputé d'un index se relit comme complet,
    # et se transmet.
    _verifier_shards(res, "search/export")

    hits = res["hits"]["hits"]
    if req.format == "docx":
        return _export_results_docx(req.query, hits)
    return _export_results_xlsx(req.query, hits)


# ── Recherches sauvegardées ─────────────────────────────────────
@app.get("/saved-searches")
def list_saved_searches(user: str = Depends(current_user)):
    username = user
    return saved_searches.list_saved(username)


@app.post("/saved-searches")
def create_saved_search(body: SavedSearchCreate, user: str = Depends(current_user)):
    username = user
    if not body.name.strip():
        raise HTTPException(status_code=400, detail="Le nom de la recherche ne peut pas être vide")
    try:
        return saved_searches.save_search(username, body.name, body.model_dump(exclude={"name"}))
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.delete("/saved-searches/{search_id}")
def remove_saved_search(search_id: str, user: str = Depends(current_user)):
    username = user
    try:
        return saved_searches.delete_saved(username, search_id)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


# ── Mon activité : historique et autocomplétion ──────────────────────
#
# Les deux lisent l'index search_logs, écrit à chaque recherche depuis
# toujours — aucune collecte nouvelle, seulement une restitution à
# l'intéressé de ce qui n'allait jusqu'ici qu'aux statistiques
# d'administration. Le détail (et les deux gisements écartés) est dans
# user_history.py.
#
# Les deux bascules démarrent à FALSE, contrairement aux plus anciennes :
# elles ajoutent un élément à l'écran, et la règle du dépôt veut qu'une
# telle bascule s'allume à la demande d'un administrateur plutôt que
# d'apparaître d'elle-même après une mise à jour.
def _require_search_history_enabled() -> None:
    if not ui_config.get_config().get("search_history_enabled", False):
        raise HTTPException(status_code=403, detail="L'historique de recherche est désactivé.")


def _require_autocomplete_enabled() -> None:
    if not ui_config.get_config().get("autocomplete_enabled", False):
        raise HTTPException(status_code=403, detail="L'autocomplétion est désactivée.")


@app.get("/me/searches")
def get_my_searches(limit: int = 10, user: str = Depends(current_user)):
    """Les dernières recherches de l'utilisateur courant.

    « De l'utilisateur courant » n'est pas un paramètre : il n'existe
    aucune route permettant de lire l'historique de quelqu'un d'autre —
    la ventilation par utilisateur, c'est /admin/search-logs, et elle
    est réservée aux administrateurs et tracée au journal d'audit.
    """
    _require_search_history_enabled()
    borne = max(1, min(limit, 50))
    try:
        return {"searches": user_history.recent_queries(es, user, borne)}
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Historique indisponible : {e}") from e


def _require_recent_documents_enabled() -> None:
    if not ui_config.get_config().get("recent_documents_enabled", False):
        raise HTTPException(status_code=403, detail="Les documents récemment consultés sont désactivés.")


@app.get("/me/recent-documents")
def get_my_recent_documents(limit: int = 10, user: str = Depends(current_user)):
    """Les derniers documents que l'utilisateur courant a ouverts.

    ⚠️ Les identifiants viennent du journal, mais les DOCUMENTS sont
    relus à travers le filtre ACL : un document dont les droits ont
    changé depuis le clic, ou qui a été supprimé de l'index, disparaît
    de cette liste. Un historique de consultation ne rouvre pas une
    porte qui s'est fermée depuis.
    """
    _require_recent_documents_enabled()
    borne = max(1, min(limit, 50))
    try:
        identifiants = user_history.recent_documents(es, user, borne)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Historique indisponible : {e}") from e
    if not identifiants:
        return {"documents": []}

    try:
        res = es.search(
            index=ES_SEARCH_ALIAS,
            size=len(identifiants),
            query={"bool": {"filter": [
                {"ids": {"values": identifiants}},
                build_acl_filter(user),
                {"terms": {"source": _searchable_source_names(user)}},
            ]}},
            source_excludes=["content", "content_vector"],
        )
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Documents indisponibles : {e}") from e

    trouves = {h["_id"]: h["_source"] for h in res["hits"]["hits"]}
    # L'ordre de consultation, pas celui d'Elasticsearch.
    return {"documents": [
        {**trouves[identifiant], "id": identifiant}
        for identifiant in identifiants
        if identifiant in trouves
    ]}


@app.get("/search/suggest")
def suggest(q: str = "", limit: int = 8, user: str = Depends(current_user)):
    """Suggestions de saisie : ses propres recherches d'abord, puis les
    auteurs, mots-clés et valeurs de facettes personnalisées du corpus
    QU'IL PEUT VOIR.

    Sous `/search/` et non à la racine, pour deux raisons : le préfixe est
    déjà proxifié par les deux Nginx et par le proxy de développement
    (comme `/search/export`), et une route `/suggest` voisinerait avec
    `/suggestions`, qui est tout autre chose — le recueil des suggestions
    d'amélioration des utilisateurs. Deux routes à une lettre d'écart sont
    un piège pour qui relira.

    Meilleur effort de bout en bout : moins de deux caractères, une panne
    d'Elasticsearch ou un dépassement du délai renvoient la liste
    (éventuellement partielle) constituée jusque-là, jamais une erreur.
    Une barre de recherche qui affiche « 503 » sous les doigts de
    l'utilisateur serait pire que pas de suggestion du tout.
    """
    _require_autocomplete_enabled()
    saisie = q.strip()
    if len(saisie) < 2:
        return {"suggestions": []}
    borne = max(1, min(limit, 20))

    propositions: list[dict] = []
    try:
        propositions += [
            {"text": entree["query"], "kind": "history", "count": entree["count"]}
            for entree in user_history.matching_queries(es, user, saisie, borne)
        ]
    except Exception as e:
        logger.warning(f"[suggest] Historique indisponible pour {user} : {e}")

    if len(propositions) < borne:
        try:
            # Exactement les filtres de /search : l'ACL de l'appelant et
            # les sources réellement cherchables. Une agrégation divulgue
            # autant qu'un résultat de recherche.
            filtres = [
                build_acl_filter(user),
                {"terms": {"source": _searchable_source_names(user)}},
            ]
            deja_vu = {p["text"].casefold() for p in propositions}
            for proposition in user_history.corpus_terms(
                es, ES_SEARCH_ALIAS, filtres, saisie, borne - len(propositions),
                champs_custom=_suggestable_custom_facets(user),
            ):
                if proposition["text"].casefold() not in deja_vu:
                    propositions.append(proposition)
        except Exception as e:
            logger.warning(f"[suggest] Corpus indisponible pour « {saisie} » : {e}")

    return {"suggestions": propositions[:borne]}


# ── Alertes sur recherches sauvegardées ──────────────────────────────
# Entièrement suspendable depuis l'admin (ui_config.alerts_enabled) :
# désactivé, toutes les routes ci-dessous renvoient 403, y compris la
# simple consultation des notifications déjà déposées — même principe
# que _require_collections_enabled() plus bas.
def _require_alerts_enabled() -> None:
    if not ui_config.get_config().get("alerts_enabled", True):
        raise HTTPException(status_code=403, detail="Les alertes sont désactivées.")


@app.patch("/saved-searches/{search_id}/alert")
def update_saved_search_alert(search_id: str, body: SavedSearchAlertUpdate, user: str = Depends(current_user)):
    _require_alerts_enabled()
    username = user
    if body.frequency not in ("daily", "weekly"):
        raise HTTPException(status_code=400, detail="frequency doit être 'daily' ou 'weekly'")
    try:
        return saved_searches.set_alert(username, search_id, body.enabled, body.frequency)
    except KeyError:
        raise HTTPException(status_code=404, detail="Recherche sauvegardée introuvable") from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.get("/alerts")
def list_alerts(user: str = Depends(current_user)):
    """Notifications in-app de l'utilisateur (voir alert_worker.py, qui
    les dépose en arrière-plan) — la plus récente en premier."""
    _require_alerts_enabled()
    username = user
    return alert_notifications.list_notifications(username)


@app.post("/alerts/{notif_id}/seen")
def mark_alert_seen(notif_id: str, user: str = Depends(current_user)):
    _require_alerts_enabled()
    username = user
    return alert_notifications.mark_seen(username, notif_id)


@app.post("/alerts/mark-all-seen")
def mark_all_alerts_seen(user: str = Depends(current_user)):
    _require_alerts_enabled()
    username = user
    return alert_notifications.mark_all_seen(username)


@app.delete("/alerts")
def purge_alerts(user: str = Depends(current_user)):
    """Efface toutes les notifications de l'utilisateur courant.

    DELETE et non POST : l'opération est une suppression, et elle est
    idempotente — la rejouer sur une liste déjà vide donne le même
    résultat.

    Chacun ne peut purger que ses propres notifications : la clé Redis
    est indexée par nom d'utilisateur résolu depuis l'en-tête X-User, il
    n'existe aucun moyen d'en désigner une autre.
    """
    _require_alerts_enabled()
    username = user
    return alert_notifications.purge(username)


# ── Collections de documents ──────────────────────────────────────────
# Strictement personnel (voir saved_collections.py) — entièrement
# suspendable depuis l'admin (ui_config.collections_enabled) : désactivé,
# toutes les routes ci-dessous renvoient 403, y compris la simple
# consultation, plutôt que de ne bloquer que la création (cohérent avec
# l'intention d'un flag "fonctionnalité désactivée" plutôt que "création
# désactivée").
def _groupes_de_partage(username: str) -> list[str]:
    """Groupes effectifs de l'utilisateur, ou liste vide si le partage de
    collections est désactivé — auquel cas les collections partagées
    cessent d'apparaître, sans que celles déjà partagées soient
    modifiées : le réglage se rallume et tout revient."""
    if not ui_config.get_config().get("collections_shared_enabled", False):
        return []
    return get_effective_groups(username)


def _require_collections_shared_enabled() -> None:
    if not ui_config.get_config().get("collections_shared_enabled", False):
        raise HTTPException(status_code=403, detail="Le partage de collections est désactivé.")


def _require_collections_enabled() -> None:
    if not ui_config.get_config().get("collections_enabled", True):
        raise HTTPException(status_code=403, detail="Les collections de documents sont désactivées.")


@app.get("/collections")
def get_collections(user: str = Depends(current_user)):
    _require_collections_enabled()
    username = user
    return saved_collections.list_collections(es, username, _groupes_de_partage(username))


@app.post("/collections")
def create_collection(body: SavedCollectionCreate, user: str = Depends(current_user)):
    _require_collections_enabled()
    username = user
    try:
        return saved_collections.create_collection(es, username, body.name)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.delete("/collections/{collection_id}")
def remove_collection(collection_id: str, user: str = Depends(current_user)):
    _require_collections_enabled()
    username = user
    try:
        return saved_collections.delete_collection(es, username, collection_id)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.post("/collections/{collection_id}/rename")
def rename_collection(collection_id: str, body: SavedCollectionRename, user: str = Depends(current_user)):
    _require_collections_enabled()
    username = user
    try:
        return saved_collections.rename_collection(es, username, collection_id, body.name)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.post("/collections/{collection_id}/documents")
def add_collection_document(collection_id: str, body: SavedCollectionDocumentAdd, user: str = Depends(current_user)):
    _require_collections_enabled()
    username = user
    source_name = _resolve_doc_source(body.doc_id)
    if source_name is not None and source_name not in _collectable_source_names():
        raise HTTPException(
            status_code=403,
            detail=f"Les documents de la source '{source_name}' ne peuvent pas être ajoutés à une collection.",
        )
    try:
        return saved_collections.add_document(es, username, collection_id, body.doc_id)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.delete("/collections/{collection_id}/documents/{doc_id}")
def remove_collection_document(collection_id: str, doc_id: str, user: str = Depends(current_user)):
    _require_collections_enabled()
    username = user
    try:
        return saved_collections.remove_document(es, username, collection_id, doc_id)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.post("/collections/{collection_id}/share")
def share_collection(
    collection_id: str, body: SavedCollectionSharing, user: str = Depends(current_user),
):
    """Partage une collection avec des groupes, ou la repasse en
    personnel (liste vide).

    ⚠️ Partager donne la RÉFÉRENCE, pas le droit de lecture : chaque
    document reste relu à travers l'ACL du lecteur. Deux personnes
    ouvrant la même collection n'y voient donc pas forcément le même
    nombre de documents, et l'interface le dit plutôt que de masquer
    l'écart.

    On ne partage qu'avec un groupe dont on est soi-même membre — sans
    cette borne, le premier usage serait de s'adresser à toute
    l'organisation.
    """
    _require_collections_enabled()
    _require_collections_shared_enabled()
    username = user
    try:
        return saved_collections.set_sharing(
            es, username, collection_id, body.groups, get_effective_groups(username),
        )
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except PermissionError as e:
        raise HTTPException(status_code=403, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.post("/collections/{collection_id}/duplicate")
def duplicate_collection(collection_id: str, user: str = Depends(current_user)):
    """Recopie dans ses propres collections une collection visible.

    C'est la porte de sortie du destinataire : il ne modifie pas la
    collection d'un autre, il s'en fait une copie."""
    _require_collections_enabled()
    username = user
    try:
        return saved_collections.duplicate_collection(
            es, username, collection_id, _groupes_de_partage(username),
        )
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


def _doc_acl(doc: dict) -> dict:
    """Normalise l'ACL d'un document, quelle que soit sa forme.

    Les documents issus de l'ingestion fichier portent un objet imbriqué
    ({"acl": {"public": true, ...}}). Une source SQL, elle, projette ses
    colonnes vers des noms de champs ES écrits en toutes lettres : le
    mapping "acl.public" produit une clé PLATE "acl.public" dans le
    _source. Elasticsearch traite les deux pareil dans une requête (c'est
    un chemin de champ), mais un accès Python par doc["acl"] ne trouvait
    rien — et tout document SQL était donc refusé par le contrôle
    ci-dessous, y compris public.

    On ne lit que les quatre clés attendues : une clé plate inconnue ne
    peut pas s'inviter dans l'ACL.
    """
    acl = dict(doc.get("acl") or {})
    for clef in ("public", "owner", "users", "groups"):
        plat = f"acl.{clef}"
        if plat in doc and clef not in acl:
            acl[clef] = doc[plat]
    return acl


def _check_doc_access(doc: dict, username: str) -> bool:
    """Même règle ACL que build_acl_filter() (public/propriétaire/partagé
    utilisateur ou groupe), mais évaluée sur un document déjà récupéré
    plutôt qu'en filtre de requête ES — utilisée partout où un document
    précis est accédé par id (GET /document, /api/preview, édition des
    mots-clés personnalisés...).

    Vérifie EN PLUS que la source du document est cherchable par cet
    utilisateur, au sens EXACT de _searchable_source_names() — le filtre
    obligatoire de /search. Un accès direct par doc_id doit respecter la
    même restriction qu'une recherche, sinon un doc_id connu (partagé,
    laissé dans une collection, deviné) la contourne.

    Trois cas étaient laissés passer jusqu'ici, tous mesurés en dev :
      - source désactivée (searchable=false) : /document renvoyait le
        contenu intégral et /api/preview servait le fichier, alors que la
        même source rendait zéro résultat en recherche ;
      - source retirée du registre : remove_source() ne supprime ni
        l'index ES ni son appartenance à l'alias, et _get_any_source()
        rendant None, le contrôle de groupes était simplement SAUTÉ —
        retirer une source ne coupait donc pas l'accès direct à ses
        documents, y compris pour une source qui était restreinte ;
      - document sans champ "source" : même raison, alors que le filtre
        {"terms": {"source": [...]}} de /search ne le remonte jamais.

    Coût : la lecture des trois registres (cache Redis de
    SOURCES_CACHE_TTL) sur une action utilisateur ponctuelle, pas dans
    une boucle de résultats.
    """
    source_name = doc.get("source")
    if not source_name or source_name not in _searchable_source_names(username):
        return False

    acl         = _doc_acl(doc)
    user_groups = get_effective_groups(username)

    return (
        acl.get("public", False)
        or acl.get("owner")  == username
        or username in acl.get("users",  [])
        or any(g in acl.get("groups", []) for g in user_groups)
    )


# ── Détail document ──────────────────────────────────────────
@app.get("/document/{doc_id}")
def get_document(
    doc_id: str,
    user: str = Depends(current_user),
):
    username = user

    doc_index = _resolve_doc_index(doc_id)
    try:
        res = es.get(index=doc_index, id=doc_id)
    except Exception as e:
        raise HTTPException(status_code=404, detail="Document introuvable") from e

    doc = res["_source"]

    if not _check_doc_access(doc, username):
        raise HTTPException(status_code=403, detail="Accès refusé")

    # Même précaution que dans /search : le champ "id" du document (une
    # source SQL y projette souvent sa clé primaire) ne doit pas passer
    # pour l'identifiant Elasticsearch, seul utilisable pour redemander
    # ce document.
    #
    # "acl" est normalisé au passage : un document SQL porte des clés
    # plates ("acl.public"), que l'affichage des droits d'accès côté
    # interface ne sait pas lire. Mieux vaut une forme unique ici que la
    # même gymnastique répétée chez chaque consommateur.
    return {**doc, "acl": _doc_acl(doc), "id": doc_id}


# ── Mots-clés personnalisés ────────────────────────────────────
# Activable/désactivable depuis l'admin (ui_config.custom_keywords_enabled)
# — désactivé, ces deux routes renvoient 403 ; les surcharges déjà
# enregistrées restent dans leur index ES (custom_keywords.py), simplement
# plus modifiables tant que le flag est désactivé (même principe que
# collections_enabled).
#
# Réservé aux documents de TYPE FICHIER ("document"/"archive_member") —
# email PST, page web, ligne SQL n'ont pas de notion de "mots-clés Office/
# PDF" à compléter.
class DocumentKeywordBody(BaseModel):
    keyword: str


def _require_custom_keywords_enabled() -> None:
    if not ui_config.get_config().get("custom_keywords_enabled", True):
        raise HTTPException(status_code=403, detail="Les mots-clés personnalisés sont désactivés.")


def _load_doc_for_keyword_edit(doc_id: str, username: str) -> tuple[str, dict]:
    """Facteur commun aux deux routes ci-dessous : résout l'index,
    récupère le document, vérifie ACL et type. Retourne (doc_index, doc)."""
    doc_index = _resolve_doc_index(doc_id)
    try:
        doc = es.get(index=doc_index, id=doc_id)["_source"]
    except Exception as e:
        raise HTTPException(status_code=404, detail="Document introuvable") from e

    if not _check_doc_access(doc, username):
        raise HTTPException(status_code=403, detail="Accès refusé")

    if doc.get("type") not in ("document", "archive_member"):
        raise HTTPException(
            status_code=400,
            detail="Les mots-clés personnalisés ne sont disponibles que pour les documents de type fichier.",
        )
    return doc_index, doc


@app.post("/document/{doc_id}/keywords")
def add_document_keyword(doc_id: str, body: DocumentKeywordBody, user: str = Depends(current_user)):
    _require_custom_keywords_enabled()
    username = user
    keyword = body.keyword.strip()
    if not keyword:
        raise HTTPException(status_code=400, detail="Mot-clé vide.")

    doc_index, doc = _load_doc_for_keyword_edit(doc_id, username)
    try:
        custom_keywords.add_keyword(es, doc_id, doc.get("source"), keyword, username)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e

    # Effet immédiat sur le document principal — la surcharge persistée
    # ci-dessus n'est réappliquée par le pipeline d'ingestion qu'à la
    # PROCHAINE réindexation (voir indexer.py:apply_keyword_overrides),
    # sans quoi l'utilisateur ne verrait pas son ajout tout de suite.
    #
    # refresh=True (pas "wait_for") : les index de documents passent par
    # restore_after_bulk() (indexer.py) après une indexation en masse, qui
    # fixe refresh_interval à 30s — "wait_for" attendrait alors jusqu'à
    # 30s le prochain rafraîchissement PLANIFIÉ au lieu d'en déclencher un
    # immédiatement. Coût négligeable ici (écriture d'un seul document,
    # action utilisateur peu fréquente), contrairement à un rafraîchissement
    # forcé pendant un bulk() de plusieurs milliers de documents.
    current = doc.get("keywords") or []
    if keyword not in current:
        current = current + [keyword]
        es.update(index=doc_index, id=doc_id, refresh=True, doc={"keywords": current})
    return {"keywords": current}


@app.delete("/document/{doc_id}/keywords/{keyword}")
def remove_document_keyword(doc_id: str, keyword: str, user: str = Depends(current_user)):
    _require_custom_keywords_enabled()
    username = user

    doc_index, doc = _load_doc_for_keyword_edit(doc_id, username)
    try:
        custom_keywords.remove_keyword(es, doc_id, doc.get("source"), keyword, username)
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e

    current = [k for k in (doc.get("keywords") or []) if k != keyword]
    if current != (doc.get("keywords") or []):
        # refresh=True — voir le commentaire équivalent dans add_document_keyword().
        es.update(index=doc_index, id=doc_id, refresh=True, doc={"keywords": current})
    return {"keywords": current}


# ── Aperçu document ──────────────────────────────────────────
@app.get("/api/preview/{doc_id}")
def preview_document(
    doc_id: str,
    user: str = Depends(current_user),
):
    # Vérification ACL via get_document (lève 403 si refusé)
    doc      = get_document(doc_id, user=user)
    filepath = doc["filepath"]
    ext      = doc["extension"]

    if "::" in filepath:
        # Document extrait d'une archive (.zip, .tar.*, .7z) — il n'existe
        # que temporairement pendant l'indexation, aucun aperçu possible.
        archive_path, member = filepath.split("::", 1)
        raise HTTPException(
            status_code=422,
            detail=f"Aperçu non disponible : document extrait de l'archive "
                   f"'{Path(archive_path).name}' (membre : {member})"
        )

    if not Path(filepath).exists():
        raise HTTPException(status_code=404, detail="Fichier introuvable")

    if ext == ".pdf":
        return FileResponse(filepath, media_type="application/pdf")
    if ext in {".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx"}:
        return _convert_to_pdf(filepath)
    if ext == ".txt":
        return FileResponse(filepath, media_type="text/plain; charset=utf-8")
    raise HTTPException(status_code=415, detail="Format non prévisualisable")


def _convert_to_pdf(filepath: str) -> StreamingResponse:
    with tempfile.TemporaryDirectory() as tmpdir:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", tmpdir, filepath],
            check=True, timeout=30
        )
        pdf_name = Path(filepath).stem + ".pdf"
        content  = Path(tmpdir, pdf_name).read_bytes()
    return StreamingResponse(
        iter([content]),
        media_type="application/pdf",
        headers={"Content-Disposition": f"inline; filename={pdf_name}"},
    )


# ── Métriques ─────────────────────────────────────────────────
@app.get("/metrics")
def get_metrics(user: str = Depends(current_user)):
    """Métriques agrégées sur TOUTES les sources (via l'alias fédéré) —
    voir /admin/status pour une ventilation par source individuelle."""
    _ensure_index_exists()
    info    = es.info()
    count   = es.count(index=ES_SEARCH_ALIAS)["count"]
    stats   = es.indices.stats(index=ES_SEARCH_ALIAS)
    size_gb = stats["_all"]["total"]["store"]["size_in_bytes"] / 1e9
    by_ext  = es.search(
        index=ES_SEARCH_ALIAS, size=0,
        aggs={"by_ext": {"terms": {"field": "extension", "size": 20}}}
    )
    return {
        "indexed":      count,
        "size_gb":      round(size_gb, 2),
        "by_extension": by_ext["aggregations"]["by_ext"]["buckets"],
        "es_version":   info["version"]["number"],
        "acl_enabled":  True,
    }


# ═══════════════════════════════════════════════════════════════
# PANNEAU D'ADMINISTRATION — /admin/*
#
# Toutes ces routes exigent Depends(require_admin) : l'utilisateur
# doit être authentifié (X-User, injecté par Nginx après validation
# SSO) ET membre du groupe ADMIN_GROUP (résolu via LDAP/AD).
#
# Aucune de ces routes n'a besoin d'un accès Docker — vérification
# d'état via le réseau applicatif (HTTP/Redis/Kafka), déclenchement
# de scan/purge via publication Kafka ou requêtes ES directes. Piloter
# le nombre de workers ou démarrer/arrêter des conteneurs reste
# réservé à `manage.sh` en CLI (voir docsearch-infra).
# ═══════════════════════════════════════════════════════════════

from fastapi import BackgroundTasks


class FiletypeUpdate(BaseModel):
    enabled: bool | None = None
    max_size_mb: float | None = None
    source: str = DEFAULT_SOURCE_NAME


class ConfigUpdate(BaseModel):
    value: str


class PathFilterPattern(BaseModel):
    pattern: str
    source: str = DEFAULT_SOURCE_NAME


class PurgeRequest(BaseModel):
    pattern: str
    source: str = DEFAULT_SOURCE_NAME
    dry_run: bool = True


class ScanRequest(BaseModel):
    source: str = DEFAULT_SOURCE_NAME
    subfolder: str | None = None


class SourceCreate(BaseModel):
    name: str
    es_index: str
    subfolder: str | None = None
    label: str | None = None
    description: str | None = None
    # None = "ne pas modifier" (distinct de False = "désactiver
    # explicitement") — nécessaire car add_source() REMPLACE l'entrée
    # entière, voir admin_add_source().
    ocr_enabled: bool | None = None


class SqlFieldMapping(BaseModel):
    column: str
    es_field: str
    es_type: str
    analyzer: str | None = None
    facet: bool = False
    facet_label: str | None = None
    # Libellé du champ dans la carte de résultat — voir FieldMapping dans
    # sql_sources_config.py. None = affiché sous un libellé dérivé,
    # "" = masqué, texte = ce libellé. La distinction None / "" doit
    # survivre à la sérialisation : ne pas remplacer par `str = ""`.
    card_label: str | None = None


class SqlSourceCreate(BaseModel):
    name: str
    db_type: str
    connection_ref: str
    query: str
    id_column: str
    es_index: str
    fields: list[SqlFieldMapping]
    poll_interval_seconds: int = sql_sources_config.DEFAULT_POLL_INTERVAL_SECONDS
    label: str | None = None
    description: str | None = None


class SqlDsnCreate(BaseModel):
    name: str
    dsn: str


class WebSourceCreate(BaseModel):
    name: str
    crawl_index: str
    es_index: str
    acl_public: bool = True
    poll_interval_seconds: int = web_sources_config.DEFAULT_POLL_INTERVAL_SECONDS
    label: str | None = None
    description: str | None = None


def _sources_status() -> dict:
    """Nombre de documents par source enregistrée — un index manquant
    (source enregistrée mais jamais indexée) compte pour 0 plutôt que de
    faire échouer tout /admin/status."""
    result = {}
    for name, source in file_sources_config.get_sources().items():
        try:
            result[name] = {
                "es_index": source.es_index,
                "label":    source.label,
                "folder":   source.folder,
                "indexed":  es.count(index=source.es_index)["count"],
            }
        except Exception:
            result[name] = {
                "es_index": source.es_index,
                "label":    source.label,
                "folder":   source.folder,
                "indexed":  0,
            }
    return result


@app.get("/admin/status")
def admin_status(user: str = Depends(require_admin)):
    """État de tous les composants : ES, Redis, Tika, Kafka, workers
    actifs, progression de l'indexation (lag), battement du watcher —
    plus une ventilation du nombre de documents par source."""
    status = cluster_status.get_full_status()
    status["sources"] = _sources_status()
    return status


class LabelUpdate(BaseModel):
    label: str


class DescriptionUpdate(BaseModel):
    description: str


class OcrUpdate(BaseModel):
    ocr_enabled: bool


@app.get("/admin/file-sources")
def admin_get_sources(user: str = Depends(require_admin)):
    return {
        name: {"es_index": s.es_index, "folder": s.folder, "label": s.label,
               "description": s.description, "ocr_enabled": s.ocr_enabled}
        for name, s in file_sources_config.get_sources().items()
    }


@app.post("/admin/file-sources")
def admin_add_source(body: SourceCreate, user: str = Depends(require_admin)):
    try:
        # add_source() REMPLACE l'entrée existante en entier — on relit
        # searchable/collectable/ocr_enabled/allowed_groups au préalable
        # pour ne pas les réinitialiser à leur valeur par défaut au
        # premier "Modifier" venu (voir add_source() docstring).
        existing = file_sources_config.get_sources().get(body.name)
        return file_sources_config.add_source(
            body.name, body.es_index, subfolder=body.subfolder, label=body.label,
            searchable=existing.searchable if existing else True,
            collectable=existing.collectable if existing else True,
            allowed_groups=list(existing.allowed_groups) if existing else None,
            description=body.description,
            ocr_enabled=body.ocr_enabled if body.ocr_enabled is not None else (existing.ocr_enabled if existing else False),
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.delete("/admin/file-sources/{name}")
def admin_remove_source(name: str, user: str = Depends(require_admin)):
    """Retire la source du registre (le watcher arrête de l'observer) —
    NE supprime PAS l'index Elasticsearch ni les documents déjà
    indexés : utiliser /admin/purge-path pour nettoyer l'existant."""
    try:
        return file_sources_config.remove_source(name)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.post("/admin/file-sources/{name}/label")
def admin_set_source_label(name: str, body: LabelUpdate, user: str = Depends(require_admin)):
    """Modifie le libellé d'affichage d'une source fichier — son nom
    (registre + champ "source" des documents déjà indexés) ne change pas."""
    try:
        return file_sources_config.set_label(name, body.label)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.post("/admin/file-sources/{name}/description")
def admin_set_source_description(name: str, body: DescriptionUpdate, user: str = Depends(require_admin)):
    try:
        return file_sources_config.set_description(name, body.description)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None


@app.post("/admin/file-sources/{name}/ocr")
def admin_set_source_ocr(name: str, body: OcrUpdate, user: str = Depends(require_admin)):
    """Active/désactive l'OCR (Tesseract via Tika) pour une source
    fichier — n'a de sens que pour les sources fichiers (PDF scannés,
    images), contrairement aux bascules génériques searchable/
    collectable : volontairement absente de
    _SOURCE_REGISTRIES//admin/all-sources/{name}/... (sql/web n'ont pas
    de notion d'OCR)."""
    try:
        return file_sources_config.set_ocr_enabled(name, body.ocr_enabled)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None


@app.get("/admin/file-sources/{name}/tree")
def admin_get_source_tree(name: str, path: str = Query(""), user: str = Depends(require_admin)):
    """Liste UN SEUL niveau de l'arborescence d'une source fichier (pour
    le chargement paresseux de la section "Arborescence des sources" côté
    admin.html) — jamais de descente récursive ici, pour rester rapide
    même sur une source à des dizaines de milliers de fichiers ; c'est au
    client de rappeler cette route à chaque dépliage de dossier.

    `path` est relatif à la racine de la source (source.folder), jamais
    un chemin absolu — mêmes conventions que path_filter (voir
    producer.py:_publish_folder pour la référence)."""
    try:
        source = file_sources_config.get_source(name)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None

    root = Path(source.folder).resolve()
    target = (root / path).resolve()
    if target != root and root not in target.parents:
        raise HTTPException(status_code=400, detail=f"Chemin invalide : '{path}' sort de la source '{name}'.")
    if not target.is_dir():
        raise HTTPException(status_code=404, detail=f"Dossier introuvable dans '{name}' : '{path}'")

    filter_config = path_filter.get_config(name)
    included_patterns = filter_config.get("included", [])
    whitelist_active = bool(included_patterns)

    entries = []
    try:
        with os.scandir(target) as it:
            for entry in it:
                # Fichiers/dossiers cachés (".git", ".DS_Store", ...) —
                # jamais indexés (voir is_excluded() côté ingestion),
                # inutile de les faire remonter ici.
                if entry.name.startswith("."):
                    continue
                rel = f"{path}/{entry.name}" if path else entry.name
                is_dir = entry.is_dir(follow_symlinks=False)
                item = {"name": entry.name, "path": rel, "type": "dir" if is_dir else "file"}
                # is_dir_excluded() ne vérifie QUE la liste noire — malgré
                # son nom (pensé pour l'élagage d'os.walk), c'est une simple
                # comparaison de motifs qui vaut aussi bien pour un fichier
                # que pour un dossier. On l'affiche pour les deux : un motif
                # noir peut viser un fichier précis (ex: "*.tmp"), pas
                # seulement un dossier entier.
                item["excluded"] = path_filter.is_dir_excluded(rel, name)
                # La liste blanche ne s'applique volontairement PAS aux
                # dossiers pour décider s'il faut les parcourir (voir
                # docstring de is_dir_excluded : "finance" ne correspond pas
                # littéralement à "finance/rapports" mais il faut quand même
                # y descendre) — ici on affiche juste, à titre informatif,
                # si CE chemin correspond explicitement à un motif inclus,
                # sans rien affirmer sur ses descendants.
                if whitelist_active:
                    item["included"] = any(path_filter.matches_pattern(rel, p) for p in included_patterns)
                entries.append(item)
    except OSError as e:
        raise HTTPException(status_code=500, detail=f"Erreur lecture dossier : {e}") from e

    entries.sort(key=lambda e: (e["type"] != "dir", e["name"].lower()))
    return {"path": path, "whitelist_active": whitelist_active, "entries": entries}


@app.get("/admin/sql-sources")
def admin_get_sql_sources(user: str = Depends(require_admin)):
    return {
        name: {
            "db_type":               s.db_type,
            "connection_ref":        s.connection_ref,
            "query":                 s.query,
            "id_column":             s.id_column,
            "es_index":              s.es_index,
            "poll_interval_seconds": s.poll_interval_seconds,
            "label":                 s.label,
            "description":           s.description,
            "fields": [
                {
                    "column": f.column, "es_field": f.es_field, "es_type": f.es_type, "analyzer": f.analyzer,
                    "facet": f.facet, "facet_label": f.facet_label,
                    # Indispensable au formulaire d'édition : ce qu'il ne
                    # LIT pas ici, il le renvoie vide au prochain
                    # enregistrement — un aller-retour par cet écran
                    # effacerait donc les libellés de carte.
                    "card_label": f.card_label,
                }
                for f in s.fields
            ],
        }
        for name, s in sql_sources_config.get_sources().items()
    }


@app.post("/admin/sql-sources")
def admin_add_sql_source(body: SqlSourceCreate, user: str = Depends(require_admin)):
    """
    Enregistre (ou met à jour) une source SQL. `connection_ref` est le
    NOM d'une variable d'environnement contenant le DSN complet — jamais
    le DSN lui-même, qui ne transite donc jamais par cette route ni par
    Redis. sql-worker (docsearch-ingestion) prend en compte la nouvelle
    source sous ~5s, sans redémarrage.
    """
    try:
        # add_source() REMPLACE l'entrée existante en entier — on relit
        # searchable/collectable/allowed_groups au préalable pour ne pas
        # les réinitialiser à leur valeur par défaut au premier "Modifier"
        # venu (voir add_source() docstring).
        existing = sql_sources_config.get_sources().get(body.name)
        return sql_sources_config.add_source(
            name=body.name,
            db_type=body.db_type,
            connection_ref=body.connection_ref,
            query=body.query,
            id_column=body.id_column,
            es_index=body.es_index,
            fields=[f.model_dump() for f in body.fields],
            poll_interval_seconds=body.poll_interval_seconds,
            label=body.label,
            searchable=existing.searchable if existing else True,
            collectable=existing.collectable if existing else True,
            allowed_groups=list(existing.allowed_groups) if existing else None,
            description=body.description,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.delete("/admin/sql-sources/{name}")
def admin_remove_sql_source(name: str, user: str = Depends(require_admin)):
    """Retire la source SQL du registre (sql-worker arrête de
    l'interroger) — NE supprime PAS l'index Elasticsearch ni les
    documents déjà indexés."""
    try:
        return sql_sources_config.remove_source(name)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None


@app.post("/admin/sql-sources/{name}/label")
def admin_set_sql_source_label(name: str, body: LabelUpdate, user: str = Depends(require_admin)):
    """Modifie le libellé d'affichage d'une source SQL — son nom
    (registre + champ "source" des documents déjà indexés) ne change pas."""
    try:
        return sql_sources_config.set_label(name, body.label)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.post("/admin/sql-sources/{name}/description")
def admin_set_sql_source_description(name: str, body: DescriptionUpdate, user: str = Depends(require_admin)):
    try:
        return sql_sources_config.set_description(name, body.description)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None


# ── DSN SQL chiffrés (registre dynamique, alternative aux variables
# d'environnement de connection_ref) — voir sql_dsn_registry.py ─────
@app.get("/admin/sql-dsns")
def admin_list_sql_dsns(user: str = Depends(require_admin)):
    """Liste les DSN dynamiques enregistrés (nom + indice non sensible
    schéma/hôte, jamais le DSN déchiffré ni son chiffré). Ne lève jamais
    (même comportement que /admin/sql-sources : Redis injoignable dégrade
    silencieusement vers une liste vide plutôt que de faire échouer la
    route)."""
    return sql_dsn_registry.list_names()


@app.post("/admin/sql-dsns")
def admin_add_sql_dsn(body: SqlDsnCreate, user: str = Depends(require_admin)):
    """
    Enregistre (ou remplace) un DSN chiffré dans Redis, sous un nom au
    format variable d'environnement — ce nom devient ensuite utilisable
    comme connection_ref d'une source SQL, à condition qu'aucune variable
    d'environnement de ce nom n'existe déjà (elle resterait sinon
    prioritaire, voir docsearch-ingestion/app/sql_indexer.py::_resolve_dsn).
    Nécessite DSN_ENCRYPTION_KEY, définie à l'identique côté docsearch-api
    (chiffrement ici) ET côté sql-worker/indexer-init (déchiffrement pour
    se connecter réellement) — voir docsearch-infra/.env.example. Aucune
    connexion à la base n'est testée ici : seule la forme du DSN est
    vérifiée.
    """
    try:
        return sql_dsn_registry.add_dsn(body.name, body.dsn)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.delete("/admin/sql-dsns/{name}")
def admin_remove_sql_dsn(name: str, user: str = Depends(require_admin)):
    """Retire un DSN chiffré du registre — toute source SQL dont le
    connection_ref pointe encore vers ce nom échouera à son prochain
    passage (sauf si une variable d'environnement du même nom existe) ;
    aucune vérification qu'une source l'utilise encore, cohérent avec
    DELETE /admin/sql-sources/{name} qui ne vérifie pas non plus les
    dépendances inverses."""
    try:
        return sql_dsn_registry.remove_dsn(name)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e


@app.get("/admin/web-sources")
def admin_get_web_sources(user: str = Depends(require_admin)):
    return {
        name: {
            "crawl_index":           s.crawl_index,
            "es_index":              s.es_index,
            "acl_public":            s.acl_public,
            "poll_interval_seconds": s.poll_interval_seconds,
            "label":                 s.label,
            "description":           s.description,
            "paused":                s.paused,
        }
        for name, s in web_sources_config.get_sources().items()
    }


@app.post("/admin/web-sources")
def admin_add_web_source(body: WebSourceCreate, user: str = Depends(require_admin)):
    """
    Enregistre (ou met à jour) une source web. `crawl_index` est l'index
    ES intermédiaire écrit par Elastic Open Web Crawler (son
    `output_index`, schéma brut du crawler) — DIFFÉRENT de `es_index`
    (schéma DocSearch final). web-worker (docsearch-ingestion) prend en
    compte la nouvelle source sous ~5s, sans redémarrage.
    """
    try:
        # add_source() REMPLACE l'entrée existante en entier — on relit
        # searchable/collectable/allowed_groups au préalable pour ne pas
        # les réinitialiser à leur valeur par défaut au premier "Modifier"
        # venu (voir add_source() docstring).
        existing = web_sources_config.get_sources().get(body.name)
        return web_sources_config.add_source(
            name=body.name,
            crawl_index=body.crawl_index,
            es_index=body.es_index,
            acl_public=body.acl_public,
            poll_interval_seconds=body.poll_interval_seconds,
            label=body.label,
            searchable=existing.searchable if existing else True,
            collectable=existing.collectable if existing else True,
            allowed_groups=list(existing.allowed_groups) if existing else None,
            description=body.description,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.delete("/admin/web-sources/{name}")
def admin_remove_web_source(name: str, user: str = Depends(require_admin)):
    """Retire la source web du registre (web-worker arrête de la
    synchroniser) — NE supprime PAS les index Elasticsearch (crawl_index
    ni es_index) ni les documents déjà indexés."""
    try:
        return web_sources_config.remove_source(name)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None


@app.post("/admin/web-sources/{name}/label")
def admin_set_web_source_label(name: str, body: LabelUpdate, user: str = Depends(require_admin)):
    """Modifie le libellé d'affichage d'une source web — son nom
    (registre + champ "source" des documents déjà indexés) ne change pas."""
    try:
        return web_sources_config.set_label(name, body.label)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.post("/admin/web-sources/{name}/description")
def admin_set_web_source_description(name: str, body: DescriptionUpdate, user: str = Depends(require_admin)):
    try:
        return web_sources_config.set_description(name, body.description)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None


class PauseUpdate(BaseModel):
    paused: bool


@app.post("/admin/web-sources/{name}/pause")
def admin_set_web_source_paused(name: str, body: PauseUpdate, user: str = Depends(require_admin)):
    """
    Suspend/reprend la synchronisation crawl_index -> es_index pour une
    source web (web-worker saute cette source à chaque tick tant que
    paused=true). Ne pilote PAS le conteneur Elastic Open Web Crawler
    lui-même (aucun accès Docker depuis cette API) : si ce conteneur
    tourne en continu (mode "schedule"), il continue d'écrire dans
    crawl_index — seule la répercussion vers DocSearch s'arrête. Les
    documents déjà indexés dans es_index restent cherchables.
    """
    try:
        return web_sources_config.set_paused(name, body.paused)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None


# ── Vue d'ensemble unifiée (les 3 types de sources confondus) ─────
# Les panneaux /admin/file-sources, /admin/sql-sources, /admin/web-sources
# ci-dessus restent le CRUD dédié à chaque type (champs spécifiques :
# dossier pour un fichier, requête pour du SQL, crawl_index pour du
# web). Cette route sert un usage différent et transverse : une seule
# liste avec le compte de documents et la bascule "recherche activée",
# indépendamment du type — pour ça, la source doit être identifiable
# sans ambiguïté par (type, name), d'où le paramètre `type` explicite
# plutôt que de chercher le nom dans les trois registres.

class SearchableUpdate(BaseModel):
    searchable: bool


class CollectableUpdate(BaseModel):
    collectable: bool


class GroupsUpdate(BaseModel):
    allowed_groups: list[str]


_SOURCE_REGISTRIES = {
    "file": file_sources_config,
    "sql":  sql_sources_config,
    "web":  web_sources_config,
}


def _all_sources_status() -> dict:
    """Fusionne les trois registres de sources en une seule liste, avec
    le nombre de documents et la taille sur disque de chaque index — un
    index manquant (source enregistrée mais jamais indexée, ou vidée)
    compte pour 0 plutôt que de faire échouer tout l'appel."""
    result = {}
    for type_, registry in _SOURCE_REGISTRIES.items():
        for name, s in registry.get_sources().items():
            try:
                indexed = es.count(index=s.es_index)["count"]
            except Exception:
                indexed = 0
            try:
                # size_in_bytes de l'index PRIMAIRE (pas x nombre de
                # replicas) — c'est l'espace occupé par les données elles-
                # mêmes, l'unité pertinente ici plutôt que l'empreinte
                # disque totale du cluster (voir /metrics pour celle-ci,
                # calculée sur l'alias fédéré ES_SEARCH_ALIAS en entier).
                size_bytes = es.indices.stats(index=s.es_index)["_all"]["primaries"]["store"]["size_in_bytes"]
            except Exception:
                size_bytes = 0
            result[name] = {
                "type":       type_,
                "es_index":   s.es_index,
                "label":       getattr(s, "label", None) or name,
                "description": getattr(s, "description", None) or "",
                "searchable":     s.searchable,
                "collectable":    s.collectable,
                "allowed_groups": list(s.allowed_groups),
                "indexed":        indexed,
                "size_bytes":     size_bytes,
            }
    return result


@app.get("/admin/all-sources")
def admin_get_all_sources(user: str = Depends(require_admin)):
    return _all_sources_status()


@app.post("/admin/all-sources/{name}/searchable")
def admin_set_source_searchable(
    name: str, body: SearchableUpdate,
    type: str = Query(..., description="file, sql ou web"),
    user: str = Depends(require_admin),
):
    """Active/désactive la RECHERCHE pour une source, quel que soit son
    type — n'affecte jamais l'ingestion (watcher/sql-worker/web-worker
    continuent normalement), seulement la visibilité dans /search."""
    registry = _SOURCE_REGISTRIES.get(type)
    if registry is None:
        raise HTTPException(status_code=400, detail=f"Type de source invalide : '{type}' (attendu file, sql ou web)")
    try:
        registry.set_searchable(name, body.searchable)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    return _all_sources_status()


@app.post("/admin/all-sources/{name}/collectable")
def admin_set_source_collectable(
    name: str, body: CollectableUpdate,
    type: str = Query(..., description="file, sql ou web"),
    user: str = Depends(require_admin),
):
    """Active/désactive l'ajout à une collection pour les documents
    d'une source, quel que soit son type — n'affecte ni l'ingestion ni
    la recherche (voir add_collection_document() et
    set_collectable() dans chaque registre)."""
    registry = _SOURCE_REGISTRIES.get(type)
    if registry is None:
        raise HTTPException(status_code=400, detail=f"Type de source invalide : '{type}' (attendu file, sql ou web)")
    try:
        registry.set_collectable(name, body.collectable)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    return _all_sources_status()


@app.post("/admin/all-sources/{name}/groups")
def admin_set_source_groups(
    name: str, body: GroupsUpdate,
    type: str = Query(..., description="file, sql ou web"),
    user: str = Depends(require_admin),
):
    """Restreint (ou lève la restriction sur) la visibilité d'une source
    dans /search aux membres d'un des groupes AD/LDAP listés — liste
    vide = aucune restriction. Quel que soit le type de source ; n'affecte
    ni l'ingestion ni l'ACL par document (voir allowed_groups dans
    file_sources_config.py pour le détail). Les noms de groupes ne sont
    PAS vérifiés contre l'annuaire LDAP (ce module n'a pas de fonction
    "lister tous les groupes" — même limite que ACCESS_GROUP/ADMIN_GROUP,
    voir auth/deps.py) : une faute de frappe rend
    silencieusement la source invisible à tout le monde plutôt que de
    lever une erreur, à l'admin de vérifier l'orthographe exacte du CN AD."""
    registry = _SOURCE_REGISTRIES.get(type)
    if registry is None:
        raise HTTPException(status_code=400, detail=f"Type de source invalide : '{type}' (attendu file, sql ou web)")
    try:
        registry.set_allowed_groups(name, body.allowed_groups)
    except KeyError as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    return _all_sources_status()


@app.get("/admin/filetypes")
def admin_get_filetypes(source: str = Query(DEFAULT_SOURCE_NAME), user: str = Depends(require_admin)):
    return filetype_config.get_config(source)


@app.post("/admin/filetypes/reset")
def admin_reset_filetypes(source: str = Query(DEFAULT_SOURCE_NAME), user: str = Depends(require_admin)):
    # Route déclarée AVANT /admin/filetypes/{extension} — sinon FastAPI
    # matcherait "reset" comme une extension et cette route ne serait
    # jamais atteinte.
    return filetype_config.reset_to_default(source)


@app.post("/admin/filetypes/{extension}")
def admin_set_filetype(extension: str, body: FiletypeUpdate, user: str = Depends(require_admin)):
    return filetype_config.set_filetype(extension, enabled=body.enabled, max_size_mb=body.max_size_mb, source=body.source)


@app.delete("/admin/filetypes/{extension}")
def admin_remove_filetype(extension: str, source: str = Query(DEFAULT_SOURCE_NAME), user: str = Depends(require_admin)):
    try:
        return filetype_config.remove_filetype(extension, source)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.get("/admin/config")
def admin_get_config(user: str = Depends(require_admin)):
    return runtime_config.get_runtime_config()


@app.post("/admin/config/reset")
def admin_reset_config(user: str = Depends(require_admin)):
    # Route déclarée AVANT /admin/config/{key} — sinon FastAPI matcherait
    # "reset" comme une clé de paramètre et cette route ne serait jamais
    # atteinte.
    return runtime_config.reset_to_default()


@app.post("/admin/config/{key}")
def admin_set_config(key: str, body: ConfigUpdate, user: str = Depends(require_admin)):
    try:
        return runtime_config.set_param(key, body.value)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None


@app.get("/admin/duplicates")
def admin_duplicates(
    source: str = Query(DEFAULT_SOURCE_NAME),
    rafraichir: bool = False,
    user: str = Depends(require_admin),
):
    """Documents indexés en plusieurs exemplaires, et place qu'ils
    occupent — par source, l'empreinte étant portée par les documents
    fichiers uniquement.

    Servi depuis un cache quotidien : sans lui, chaque ouverture du
    panneau lancerait une agrégation sur tout l'index pendant que les
    utilisateurs cherchent. `rafraichir=true` force le recalcul.
    """
    s = _get_any_source(source)
    try:
        return duplicates.rapport(es, s.es_index, rafraichir=rafraichir)
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Rapport indisponible : {e}") from e


# ── Thésaurus (synonymes de recherche) ───────────────────────────────
class SynonymRule(BaseModel):
    regle: str


class SynonymTest(BaseModel):
    texte: str
    source: str = DEFAULT_SOURCE_NAME


@app.get("/admin/synonyms")
def admin_get_synonyms(user: str = Depends(require_admin)):
    try:
        return {"regles": synonyms.lister(es), "jeu": synonyms.SYNONYMS_SET}
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Thésaurus indisponible : {e}") from e


@app.post("/admin/synonyms")
def admin_add_synonym(body: SynonymRule, user: str = Depends(require_admin)):
    """Ajoute ou remplace une règle. Effet immédiat : Elasticsearch
    recharge lui-même les analyseurs des index concernés, sans
    réindexation — le nombre de shards rechargés est remonté, c'est la
    seule preuve que la modification est en vigueur."""
    try:
        return synonyms.ajouter(es, body.regle)
    except synonyms.RegleInvalide as e:
        raise HTTPException(status_code=400, detail=str(e)) from None
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Écriture impossible : {e}") from e


@app.delete("/admin/synonyms/{rule_id}")
def admin_remove_synonym(rule_id: str, user: str = Depends(require_admin)):
    try:
        return synonyms.supprimer(es, rule_id)
    except synonyms.RegleInvalide as e:
        raise HTTPException(status_code=404, detail=str(e)) from None
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Suppression impossible : {e}") from e


@app.post("/admin/synonyms/test")
def admin_test_synonyms(body: SynonymTest, user: str = Depends(require_admin)):
    """Ce que le moteur comprend d'une requête, synonymes appliqués.

    Indispensable : une règle mal écrite ne produit aucune erreur,
    seulement une recherche qui ne trouve rien de plus qu'avant."""
    s = _get_any_source(body.source)
    try:
        return synonyms.tester(es, s.es_index, body.texte)
    except Exception as e:
        raise HTTPException(
            status_code=503,
            detail=(
                f"Analyse impossible : {e}. L'index de cette source a-t-il reçu "
                "l'analyseur de synonymes (./manage.sh migrer-synonymes) ?"
            ),
        ) from e


# ── Résultats épinglés ───────────────────────────────────────────────
class PinnedRule(BaseModel):
    requete: str
    documents: list[str] = []


@app.get("/admin/pinned")
def admin_get_pinned(user: str = Depends(require_admin)):
    """Le registre, enrichi de l'état de chaque document épinglé.

    Un identifiant qui ne correspond plus à rien (document supprimé,
    source retirée) est signalé : sans ça, on épingle durablement un lien
    mort que personne ne voit disparaître — l'utilisateur, lui, ne voit
    rien du tout, la relecture le filtrant."""
    regles = pinned.lister()
    identifiants = [d for regle in regles for d in regle["documents"]]
    existants: dict[str, dict] = {}
    if identifiants:
        try:
            res = es.search(
                index=ES_SEARCH_ALIAS,
                size=len(identifiants),
                query={"ids": {"values": identifiants}},
                source_includes=["filename", "title", "filepath", "source"],
            )
            existants = {h["_id"]: h["_source"] for h in res["hits"]["hits"]}
        except Exception as e:
            logger.warning(f"[epingles] État des documents indisponible : {e}")

    return {"regles": [
        {
            "requete": regle["requete"],
            "documents": [
                {"id": identifiant, "trouve": identifiant in existants, **existants.get(identifiant, {})}
                for identifiant in regle["documents"]
            ],
        }
        for regle in regles
    ]}


@app.post("/admin/pinned")
def admin_set_pinned(body: PinnedRule, user: str = Depends(require_admin)):
    """Remplace les épinglages d'une requête. Une liste vide les retire."""
    try:
        pinned.definir(body.requete, body.documents)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    return admin_get_pinned(user=user)


@app.get("/admin/retention")
def admin_get_retention(user: str = Depends(require_admin)):
    """Ce que la purge quotidienne des journaux emporterait, journal par
    journal — sans rien supprimer.

    Les durées elles-mêmes se règlent dans les paramètres opérationnels
    (`retention_*_days`, voir `/admin/config`), comme le reste : cette
    route ne fait que montrer leur effet. Un réglage destructeur qu'on ne
    peut pas prévisualiser ne se règle jamais, ou se règle une fois de
    trop.
    """
    return {"journaux": log_retention.apercu(es)}


@app.get("/admin/path-filters")
def admin_get_path_filters(source: str = Query(DEFAULT_SOURCE_NAME), user: str = Depends(require_admin)):
    return path_filter.get_config(source)


@app.post("/admin/path-filters/exclude")
def admin_exclude_path(body: PathFilterPattern, user: str = Depends(require_admin)):
    return path_filter.add_excluded(body.pattern, body.source)


@app.post("/admin/path-filters/include")
def admin_include_path(body: PathFilterPattern, user: str = Depends(require_admin)):
    return path_filter.add_included(body.pattern, body.source)


@app.post("/admin/path-filters/remove")
def admin_remove_path_filter(body: PathFilterPattern, user: str = Depends(require_admin)):
    # POST plutôt que DELETE avec le motif dans l'URL : un motif comme
    # "finance/confidentiel" contient des "/" qui casseraient un
    # paramètre de chemin FastAPI.
    return path_filter.remove_filter(body.pattern, body.source)


@app.post("/admin/purge-path")
def admin_purge_path(body: PurgeRequest, user: str = Depends(require_admin)):
    """dry_run=True (défaut) : aperçu sans suppression. Toujours
    appeler en dry-run d'abord depuis l'interface avant confirmation.
    Opère sur l'index de `body.source` uniquement (défaut : source
    par défaut, rétrocompatible avec un client qui n'envoie pas ce champ)."""
    try:
        n = admin_scan.purge_path(body.pattern, source_name=body.source, dry_run=body.dry_run)
        return {"pattern": body.pattern, "source": body.source, "dry_run": body.dry_run, "matched": n}
    except KeyError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e)) from e


@app.post("/admin/scan")
def admin_trigger_scan(
    body: ScanRequest,
    background_tasks: BackgroundTasks,
    user: str = Depends(require_admin),
):
    """
    Déclenche un scan (publication Kafka) en arrière-plan pour UNE
    source — ne bloque pas la requête HTTP le temps de parcourir tout
    son dossier. Suivre la progression via GET /admin/status
    (workers.pending_documents).
    """
    def _run():
        try:
            result = admin_scan.trigger_scan(body.source, body.subfolder)
            logger.info(f"[admin] Scan terminé par {user} : {result}")
        except Exception as e:
            logger.error(f"[admin] Scan déclenché par {user} a échoué : {e}")

    background_tasks.add_task(_run)
    return {"status": "démarré", "source": body.source, "subfolder": body.subfolder or "(dossier complet)"}


# ── Mesure de satisfaction (pouce, NPS, clics, suggestions) ─────
# Quatre signaux distincts, volontairement pas fusionnés :
#   - feedback (pouce haut/bas) : explicite, par recherche (search_id).
#   - NPS : explicite, sur l'outil en général, PAS rattaché à une
#     recherche précise — occasionnel (cadence gérée côté client).
#   - clics : implicite, toujours actif (aucun flag), par recherche.
#   - suggestions : explicite, texte libre, PAS rattaché à une recherche
#     précise — point d'entrée permanent dans l'en-tête (index.html).
# feedback/NPS/suggestions sont individuellement suspendables
# (engagement_config.py) sans redémarrage ; le tracking de clic n'a pas
# cette option (signal passif, aucune UI ni friction ajoutée).

class FeedbackCreate(BaseModel):
    search_id: str
    rating: str  # "up" | "down"


class ClickCreate(BaseModel):
    search_id: str
    doc_id: str
    position: int


class NpsCreate(BaseModel):
    score: int = Field(ge=0, le=10)


class SuggestionCreate(BaseModel):
    text: str
    category: str | None = None   # "bug" | "idea" | "other", libre (pas de contrainte serveur)
    anonymous: bool = True        # défaut anonyme — l'utilisateur doit explicitement décocher pour être identifié


class EngagementConfigUpdate(BaseModel):
    feedback_enabled:    bool | None = None
    nps_enabled:         bool | None = None
    suggestions_enabled: bool | None = None


# ── Bascules d'interface (distinct de la mesure de satisfaction) ──
class UiConfigUpdate(BaseModel):
    chat_enabled:        bool | None = None
    footer_enabled:      bool | None = None
    footer_enabled_admin: bool | None = None
    admin_links_enabled: bool | None = None
    export_enabled:      bool | None = None
    help_enabled:        bool | None = None
    collections_enabled: bool | None = None
    custom_keywords_enabled: bool | None = None
    alerts_enabled:      bool | None = None
    search_history_enabled: bool | None = None
    autocomplete_enabled: bool | None = None
    recent_documents_enabled: bool | None = None
    collections_shared_enabled: bool | None = None
    sort_enabled:        bool | None = None
    search_time_enabled: bool | None = None
    acl_visible_enabled: bool | None = None
    shortcuts_link_enabled: bool | None = None
    empty_state_animation_enabled: bool | None = None
    show_current_user_enabled: bool | None = None
    show_current_user_groups_enabled: bool | None = None
    show_current_user_enabled_admin: bool | None = None
    show_current_user_groups_enabled_admin: bool | None = None
    theme: str | None = None
    theme_admin: str | None = None
    header_logo_url: str | None = None
    logo_text: str | None = None
    header_logo_text: str | None = None
    header_subtitle_text: str | None = None
    favicon_url: str | None = None
    footer_text: str | None = None
    footer_bottom_text: str | None = None
    # Page de connexion — voir ui_config.py pour le détail (URL vide =
    # lien masqué, ProConnect = jalon désactivé).
    login_inscription_url: str | None = None
    login_mot_de_passe_oublie_url: str | None = None
    login_proconnect_enabled: bool | None = None
    sources_mount_display: str | None = None


@app.get("/ui-config")
def get_ui_config():
    """**Volontairement PUBLIQUE**, avec /health et /is-admin — et ce
    n'est pas un reliquat : la page de connexion, qui est justement celle
    qu'on atteint SANS session, l'appelle pour son bloc-marque et son
    titre. L'exiger authentifiée donnerait un écran de connexion sans
    identité visuelle. Elle ne porte que des bascules d'affichage.

    L'interface de recherche l'appelle pour savoir si le lien "Assistant
    IA" doit être affiché dans l'en-tête.

    Ajoute "sources_mount" (préfixe réel des chemins stockés dans ES,
    ex: "/sources") en lecture seule à côté des champs persistés dans
    Redis — c'est une variable d'environnement (SOURCES_MOUNT, voir
    file_sources_config.py), pas un réglage admin : index.html s'en
    sert pour savoir quel préfixe remplacer par "sources_mount_display"
    dans copyPathClick()."""
    config = dict(ui_config.get_config())
    config["sources_mount"] = file_sources_config.SOURCES_MOUNT
    return config


@app.get("/is-admin")
def get_is_admin(user: str | None = Depends(optional_user)):
    """
    Route PUBLIQUE, et la seule à l'être avec /health : elle ne renvoie
    jamais 401/403. L'interface l'appelle pour savoir si les liens
    "Administration"/"Statistiques" doivent être affichés — une page qui
    échouerait de toute façon en 403 n'a pas à être proposée — et pour
    afficher l'identité connectée dans l'en-tête.

    D'où `optional_user` et non `current_user` : un visiteur sans session
    doit recevoir `{"is_admin": false, "user": null}` et être renvoyé vers
    la page de connexion par l'interface, pas se voir opposer un 401 que
    l'en-tête ne saurait pas afficher.

    `groups` sont les groupes EFFECTIFS (annuaire ∪ compte de secours),
    liste vide si personne n'est authentifié.
    """
    return {
        "is_admin": is_admin(user),
        "user": user,
        "groups": get_effective_groups(user) if user else [],
    }


@app.post("/admin/ui-config")
def admin_set_ui_config(body: UiConfigUpdate, user: str = Depends(require_admin)):
    """Active/désactive des éléments d'interface (ex: lien Assistant IA,
    pied de page) — effectif immédiatement pour toute nouvelle page
    chargée."""
    try:
        config = ui_config.get_config()
        if body.chat_enabled is not None:
            config = ui_config.set_param("chat_enabled", body.chat_enabled)
        if body.footer_enabled is not None:
            config = ui_config.set_param("footer_enabled", body.footer_enabled)
        if body.footer_enabled_admin is not None:
            config = ui_config.set_param("footer_enabled_admin", body.footer_enabled_admin)
        if body.admin_links_enabled is not None:
            config = ui_config.set_param("admin_links_enabled", body.admin_links_enabled)
        if body.export_enabled is not None:
            config = ui_config.set_param("export_enabled", body.export_enabled)
        if body.help_enabled is not None:
            config = ui_config.set_param("help_enabled", body.help_enabled)
        if body.collections_enabled is not None:
            config = ui_config.set_param("collections_enabled", body.collections_enabled)
        if body.custom_keywords_enabled is not None:
            config = ui_config.set_param("custom_keywords_enabled", body.custom_keywords_enabled)
        if body.alerts_enabled is not None:
            config = ui_config.set_param("alerts_enabled", body.alerts_enabled)
        if body.search_history_enabled is not None:
            config = ui_config.set_param("search_history_enabled", body.search_history_enabled)
        if body.autocomplete_enabled is not None:
            config = ui_config.set_param("autocomplete_enabled", body.autocomplete_enabled)
        if body.recent_documents_enabled is not None:
            config = ui_config.set_param("recent_documents_enabled", body.recent_documents_enabled)
        if body.collections_shared_enabled is not None:
            config = ui_config.set_param("collections_shared_enabled", body.collections_shared_enabled)
        if body.sort_enabled is not None:
            config = ui_config.set_param("sort_enabled", body.sort_enabled)
        if body.search_time_enabled is not None:
            config = ui_config.set_param("search_time_enabled", body.search_time_enabled)
        if body.acl_visible_enabled is not None:
            config = ui_config.set_param("acl_visible_enabled", body.acl_visible_enabled)
        if body.shortcuts_link_enabled is not None:
            config = ui_config.set_param("shortcuts_link_enabled", body.shortcuts_link_enabled)
        if body.empty_state_animation_enabled is not None:
            config = ui_config.set_param("empty_state_animation_enabled", body.empty_state_animation_enabled)
        if body.show_current_user_enabled is not None:
            config = ui_config.set_param("show_current_user_enabled", body.show_current_user_enabled)
        if body.show_current_user_groups_enabled is not None:
            config = ui_config.set_param("show_current_user_groups_enabled", body.show_current_user_groups_enabled)
        if body.show_current_user_enabled_admin is not None:
            config = ui_config.set_param("show_current_user_enabled_admin", body.show_current_user_enabled_admin)
        if body.show_current_user_groups_enabled_admin is not None:
            config = ui_config.set_param("show_current_user_groups_enabled_admin", body.show_current_user_groups_enabled_admin)
        if body.theme is not None:
            config = ui_config.set_theme(body.theme, "theme")
        if body.theme_admin is not None:
            config = ui_config.set_theme(body.theme_admin, "theme_admin")
        if body.header_logo_url is not None:
            config = ui_config.set_text("header_logo_url", body.header_logo_url)
        if body.logo_text is not None:
            config = ui_config.set_text("logo_text", body.logo_text)
        if body.header_logo_text is not None:
            config = ui_config.set_text("header_logo_text", body.header_logo_text)
        if body.header_subtitle_text is not None:
            config = ui_config.set_text("header_subtitle_text", body.header_subtitle_text)
        if body.favicon_url is not None:
            config = ui_config.set_text("favicon_url", body.favicon_url)
        if body.footer_text is not None:
            config = ui_config.set_text("footer_text", body.footer_text)
        if body.footer_bottom_text is not None:
            config = ui_config.set_text("footer_bottom_text", body.footer_bottom_text)
        if body.login_inscription_url is not None:
            config = ui_config.set_text("login_inscription_url", body.login_inscription_url)
        if body.login_mot_de_passe_oublie_url is not None:
            config = ui_config.set_text("login_mot_de_passe_oublie_url", body.login_mot_de_passe_oublie_url)
        if body.login_proconnect_enabled is not None:
            config = ui_config.set_param("login_proconnect_enabled", body.login_proconnect_enabled)
        if body.sources_mount_display is not None:
            config = ui_config.set_text("sources_mount_display", body.sources_mount_display)
        return config
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.get("/engagement-config")
def get_engagement_config(user: str = Depends(current_user)):
    """
    Exige une session, aucun droit particulier — l'interface de recherche l'appelle pour savoir
    si le pouce et le NPS doivent être affichés. Ne PAS confondre avec
    /admin/engagement-config (même donnée, réservé à l'admin pour
    modification) : cette route-ci n'expose rien de sensible.
    """
    return engagement_config.get_config()


@app.post("/feedback")
def submit_feedback(body: FeedbackCreate, request: Request, user: str = Depends(current_user)):
    """
    Enregistre un pouce haut/bas pour une recherche précise (search_id
    renvoyé par POST /search). Simple mise à jour partielle du document
    search_logs déjà existant — écrase un avis précédent sur la même
    recherche plutôt que d'en accumuler plusieurs (un seul avis a du sens
    par recherche).
    """
    if not engagement_config.get_config()["feedback_enabled"]:
        raise HTTPException(status_code=403, detail="Le recueil d'avis est désactivé.")
    if body.rating not in ("up", "down"):
        raise HTTPException(status_code=400, detail="rating doit être 'up' ou 'down'.")
    try:
        es.update(index=search_log.SEARCH_LOG_INDEX, id=body.search_id, doc={"feedback": body.rating})
    except Exception as e:
        if "not_found" in str(e).lower():
            raise HTTPException(status_code=404, detail="search_id introuvable.") from None
        raise HTTPException(status_code=500, detail=str(e)) from e
    return {"status": "ok"}


@app.post("/click")
def submit_click(body: ClickCreate, user: str = Depends(current_user)):
    """
    Enregistre le clic sur UN résultat d'une recherche précise (position
    dans la liste, 0-indexée) — signal toujours actif, pas de flag
    d'activation (voir docstring de section). Append via script Painless
    plutôt qu'une mise à jour de champ simple : "clicks" est une LISTE,
    un même search_id peut recevoir plusieurs clics (résultats consultés
    un par un avant de trouver le bon).
    """
    try:
        es.update(
            index=search_log.SEARCH_LOG_INDEX,
            id=body.search_id,
            script={
                "source": (
                    "if (ctx._source.clicks == null) { ctx._source.clicks = [] } "
                    "ctx._source.clicks.add(params.click)"
                ),
                "params": {
                    "click": {
                        "doc_id":    body.doc_id,
                        "position":  body.position,
                        "timestamp": datetime.now(timezone.utc).isoformat(),
                    }
                },
            },
        )
    except Exception as e:
        # Best-effort : un clic non enregistré (search_id déjà expiré,
        # ES momentanément indisponible...) ne doit jamais remonter comme
        # erreur visible — l'utilisateur est en train de consulter un
        # document, pas d'interagir avec le tracking.
        logger.warning(f"[click] Échec d'enregistrement pour search_id={body.search_id} : {e}")
    return {"status": "ok"}


@app.post("/nps")
def submit_nps(body: NpsCreate, user: str = Depends(current_user)):
    """Enregistre une réponse NPS (0-10), indépendamment de toute
    recherche précise — voir nps_log.py."""
    if not engagement_config.get_config()["nps_enabled"]:
        raise HTTPException(status_code=403, detail="Le NPS est désactivé.")
    username = user
    nps_log.log_nps(es, username=username, score=body.score, groups=get_effective_groups(username))
    return {"status": "ok"}


@app.post("/suggestions")
def submit_suggestion(body: SuggestionCreate, user: str = Depends(current_user)):
    """Enregistre une suggestion libre, indépendamment de toute recherche
    précise — voir suggestion_log.py. Anonyme par défaut ; l'identité
    n'est résolue via X-User que si l'utilisateur a explicitement décoché
    "rester anonyme" côté UI (body.anonymous == False)."""
    if not engagement_config.get_config()["suggestions_enabled"]:
        raise HTTPException(status_code=403, detail="Le recueil de suggestions est désactivé.")
    text = body.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="La suggestion ne peut pas être vide.")
    username = None if body.anonymous else user
    # Groupes seulement si l'utilisateur s'est identifié : une suggestion
    # anonyme ne doit rien porter qui permette de la rattacher à un
    # service (voir suggestion_log.py).
    suggestion_log.log_suggestion(
        es,
        text=text,
        category=body.category,
        username=username,
        groups=get_effective_groups(username) if username else None,
    )
    return {"status": "ok"}


@app.post("/admin/engagement-config")
def admin_set_engagement_config(body: EngagementConfigUpdate, user: str = Depends(require_admin)):
    """Active/désactive le pouce, le NPS et/ou les suggestions —
    effectif immédiatement pour toute nouvelle page chargée (l'UI relit
    /engagement-config à chaque chargement, pas de cache long côté
    client)."""
    try:
        config = engagement_config.get_config()
        if body.feedback_enabled is not None:
            config = engagement_config.set_param("feedback_enabled", body.feedback_enabled)
        if body.nps_enabled is not None:
            config = engagement_config.set_param("nps_enabled", body.nps_enabled)
        if body.suggestions_enabled is not None:
            config = engagement_config.set_param("suggestions_enabled", body.suggestions_enabled)
        return config
    except (ValueError, RuntimeError) as e:
        raise HTTPException(status_code=400, detail=str(e)) from e


@app.get("/admin/nps-summary")
def admin_nps_summary(user: str = Depends(require_admin)):
    """Score NPS agrégé + répartition détracteurs/passifs/promoteurs,
    pour la page /stats.html."""
    return nps_log.summary(es)


@app.get("/admin/suggestions")
def admin_list_suggestions(
    user:  str = Depends(require_admin),
    size:  int = 50,
    from_: int = Query(0, alias="from"),
):
    """Liste paginée des suggestions, plus récentes d'abord — pour la
    page /stats.html."""
    return suggestion_log.list_suggestions(es, size=size, from_=from_)


class SuggestionStatusUpdate(BaseModel):
    status: str


@app.post("/admin/suggestions/{suggestion_id}/status")
def admin_set_suggestion_status(suggestion_id: str, body: SuggestionStatusUpdate, user: str = Depends(require_admin)):
    """Suivi de traitement d'une suggestion (nouveau/en_cours/traite) —
    purement interne à l'équipe, n'informe jamais l'auteur (voir
    suggestion_log.py : l'anonymat par défaut rend une notification
    impossible à garantir de toute façon)."""
    try:
        suggestion_log.set_status(es, suggestion_id=suggestion_id, status=body.status)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from None
    except Exception as e:
        raise HTTPException(status_code=404, detail=f"Suggestion introuvable : {e}") from e
    return {"status": "ok"}


# ── Statistiques de recherche ───────────────────────────────────
# Reprend les clés de la section "timing" du résumé pour le cas où
# l'index n'existe pas encore (première installation) — la page de stats
# affiche alors des tirets plutôt que de tomber sur une clé absente.
_EMPTY_TIMING_SUMMARY = {
    "avg_ms": None, "p50_ms": None, "p95_ms": None, "took_avg_ms": None,
    "slow_count": 0, "slow_threshold_ms": SLOW_SEARCH_MS, "measured": 0,
}


def _round_ms(value: float | None) -> float | None:
    """Arrondit une durée au dixième de milliseconde. None reste None :
    une agrégation sur un index où aucune recherche n'a encore été
    mesurée ne renvoie pas 0, elle ne renvoie rien — et afficher « 0 ms »
    au lieu de « — » ferait croire à une recherche instantanée."""
    return None if value is None else round(value, 1)


@app.get("/admin/search-logs/summary")
def admin_search_logs_summary(user: str = Depends(require_admin)):
    """Compteurs agrégés + répartition par jour (14 derniers jours) +
    temps de recherche pour les cartes de résumé de la page /stats.html."""
    try:
        res = es.search(
            index=search_log.SEARCH_LOG_INDEX,
            size=0,
            aggs={
                "unique_users": {"cardinality": {"field": "username"}},
                "unique_ips":   {"cardinality": {"field": "ip"}},
                "by_day": {
                    "date_histogram": {"field": "timestamp", "calendar_interval": "day"},
                },
                "feedback_up":   {"filter": {"term": {"feedback": "up"}}},
                "feedback_down": {"filter": {"term": {"feedback": "down"}}},
                # Avis par groupe. `missing` donne un lot explicite aux
                # documents sans groupe — historique d'avant la capture,
                # ou utilisateur sans appartenance : sans lui, la somme
                # des lots ne retomberait pas sur le total global et
                # ferait douter de tout le tableau.
                "by_group": {
                    "terms": {"field": "groups", "size": 50, "missing": "__sans_groupe__"},
                    "aggs": {
                        "feedback_up":   {"filter": {"term": {"feedback": "up"}}},
                        "feedback_down": {"filter": {"term": {"feedback": "down"}}},
                    },
                },
                # Temps de recherche. `duration_measured` n'est pas une
                # redondance de total_searches : les recherches d'avant
                # l'introduction de la mesure n'ont pas le champ et sont
                # ignorées par avg/percentiles. Sans ce décompte, une
                # moyenne calculée sur douze lignes se lirait comme
                # portant sur tout l'historique.
                "duration_avg":         {"avg": {"field": "duration_ms"}},
                "duration_percentiles": {"percentiles": {"field": "duration_ms", "percents": [50, 95]}},
                "duration_measured":    {"value_count": {"field": "duration_ms"}},
                "took_avg":             {"avg": {"field": "took_ms"}},
                "slow_searches": {
                    "filter": {"range": {"duration_ms": {"gte": SLOW_SEARCH_MS}}},
                } if SLOW_SEARCH_MS else {"filter": {"match_none": {}}},
            },
        )
    except Exception as e:
        if "index_not_found" in str(e).lower():
            return {"total_searches": 0, "unique_users": 0, "unique_ips": 0, "by_day": [],
                     "feedback_up": 0, "feedback_down": 0,
                     "by_group": [], "searches_by_group": [],
                     "timing": _EMPTY_TIMING_SUMMARY}
        raise HTTPException(status_code=500, detail=str(e)) from e

    percentiles = res["aggregations"]["duration_percentiles"]["values"]

    return {
        "total_searches": res["hits"]["total"]["value"],
        "unique_users":    res["aggregations"]["unique_users"]["value"],
        "unique_ips":      res["aggregations"]["unique_ips"]["value"],
        "by_day": [
            {"date": b["key_as_string"][:10], "count": b["doc_count"]}
            for b in res["aggregations"]["by_day"]["buckets"][-14:]
        ],
        "feedback_up":   res["aggregations"]["feedback_up"]["doc_count"],
        "feedback_down": res["aggregations"]["feedback_down"]["doc_count"],
        # Volume de recherches par groupe. Même agrégation que ci-dessous
        # — inutile d'en lancer une seconde — mais TOUS les lots sont
        # gardés : un groupe qui cherche beaucoup sans jamais donner son
        # avis a précisément sa place ici, alors qu'il n'aurait rien à
        # dire dans un tableau d'avis.
        "searches_by_group": [
            {"group": b["key"], "count": b["doc_count"]}
            for b in res["aggregations"]["by_group"]["buckets"]
        ],
        # Un utilisateur de deux groupes compte dans les deux : la somme
        # des lots dépasse donc le total global. C'est le propre d'une
        # agrégation par groupe, et c'est écrit sur la page.
        "by_group": [
            {
                "group":         b["key"],
                "searches":      b["doc_count"],
                "feedback_up":   b["feedback_up"]["doc_count"],
                "feedback_down": b["feedback_down"]["doc_count"],
            }
            for b in res["aggregations"]["by_group"]["buckets"]
            # Les groupes sans le moindre avis n'apprennent rien sur la
            # satisfaction : ils encombreraient un tableau qui ne parle
            # que de ça.
            if b["feedback_up"]["doc_count"] or b["feedback_down"]["doc_count"]
        ],
        "timing": {
            "avg_ms":            _round_ms(res["aggregations"]["duration_avg"]["value"]),
            # Les clés d'une agrégation "percentiles" sont les percentiles
            # eux-mêmes, en flottant ("50.0"), pas les entiers demandés.
            "p50_ms":            _round_ms(percentiles.get("50.0")),
            "p95_ms":            _round_ms(percentiles.get("95.0")),
            "took_avg_ms":       _round_ms(res["aggregations"]["took_avg"]["value"]),
            "slow_count":        res["aggregations"]["slow_searches"]["doc_count"],
            "slow_threshold_ms": SLOW_SEARCH_MS,
            "measured":          res["aggregations"]["duration_measured"]["value"],
        },
    }


@app.get("/admin/search-logs/zero-results")
def admin_zero_result_searches(user: str = Depends(require_admin), size: int = 50):
    """Requêtes ayant retourné 0 résultat, groupées et comptées (les plus
    fréquentes en premier) — à partir des logs déjà collectés par chaque
    recherche (voir search_log.py), aucun nouveau tracking nécessaire.
    Aide à repérer du contenu manquant ou des requêtes mal formulées."""
    try:
        res = es.search(
            index=search_log.SEARCH_LOG_INDEX,
            size=0,
            query={"term": {"total_results": 0}},
            aggs={
                "by_query": {
                    "terms": {"field": "query.keyword", "size": size, "order": {"_count": "desc"}},
                    "aggs": {
                        "last_seen": {"max": {"field": "timestamp", "format": "strict_date_optional_time"}},
                    },
                },
                # Aucune capture supplémentaire n'est nécessaire : le champ
                # `groups` est déjà écrit à chaque recherche.
                "by_group": {
                    "terms": {"field": "groups", "size": 50, "missing": "__sans_groupe__"},
                },
            },
        )
    except Exception as e:
        if "index_not_found" in str(e).lower():
            return {"total_zero_result_searches": 0, "results": [], "by_group": []}
        raise HTTPException(status_code=500, detail=str(e)) from e

    return {
        "total_zero_result_searches": res["hits"]["total"]["value"],
        "results": [
            {
                "query":     b["key"],
                "count":     b["doc_count"],
                "last_seen": b["last_seen"]["value_as_string"],
            }
            for b in res["aggregations"]["by_query"]["buckets"]
        ],
        "by_group": [
            {"group": b["key"], "count": b["doc_count"]}
            for b in res["aggregations"]["by_group"]["buckets"]
        ],
    }


# ── Journal d'audit ──────────────────────────────────────────────
@app.get("/admin/audit-log")
def admin_get_audit_log(
    user:  str = Depends(require_admin),
    size:  int = 50,
    from_: int = Query(0, alias="from"),
):
    """Liste paginée des actions d'administration, plus récentes
    d'abord — alimentée par audit_log_middleware, voir audit_log.py."""
    return audit_log.list_actions(es, size=size, from_=from_)


def _search_logs_query(q: str | None, username: str | None) -> dict:
    """Filtre partagé entre /admin/search-logs (paginé) et
    /admin/search-logs/export (export complet) — mêmes critères."""
    must = []
    if q:
        must.append({"match": {"query": q}})
    if username:
        must.append({"term": {"username": username.lower()}})
    return {"bool": {"must": must}} if must else {"match_all": {}}


@app.get("/admin/search-logs")
def admin_search_logs(
    user:     str = Depends(require_admin),
    q:        str | None = None,
    username: str | None = None,
    size:     int = 50,
    from_:    int = Query(0, alias="from"),
):
    """Liste paginée des recherches effectuées, plus récentes d'abord —
    qui, quand, depuis quelle IP, quelle requête, combien de résultats."""
    query = _search_logs_query(q, username)

    try:
        res = es.search(
            index=search_log.SEARCH_LOG_INDEX,
            query=query,
            sort=[{"timestamp": {"order": "desc"}}],
            size=size,
            from_=from_,
        )
    except Exception as e:
        if "index_not_found" in str(e).lower():
            return {"total": 0, "results": []}
        raise HTTPException(status_code=500, detail=str(e)) from e

    return {
        "total":   res["hits"]["total"]["value"],
        # Même précaution que dans /search : l'identifiant ES doit
        # l'emporter sur un éventuel champ "id" du document.
        "results": [{**h["_source"], "id": h["_id"]} for h in res["hits"]["hits"]],
    }


# Plafond de lignes exportées : au-delà, l'export reste utilisable (les
# N premières lignes, plus récentes d'abord) plutôt que de saturer la
# mémoire de l'API ou du navigateur sur un historique de plusieurs
# centaines de milliers de recherches.
SEARCH_LOGS_EXPORT_MAX_ROWS = 20_000


def _join(value) -> str:
    """Aplati une valeur potentiellement multi-valuée (extension, author,
    source...) en texte lisible dans une cellule de tableur."""
    if value is None:
        return ""
    if isinstance(value, list):
        return ", ".join(str(v) for v in value)
    return str(value)


@app.get("/admin/search-logs/export")
def admin_export_search_logs(
    user:     str = Depends(require_admin),
    q:        str | None = None,
    username: str | None = None,
):
    """
    Export XLSX de l'historique des recherches — mêmes filtres que
    GET /admin/search-logs (q, username), mais TOUTES les lignes
    correspondantes (jusqu'à SEARCH_LOGS_EXPORT_MAX_ROWS) plutôt qu'une
    seule page, pour analyse hors-ligne dans un tableur.
    """
    from openpyxl import Workbook

    query = _search_logs_query(q, username)

    wb = Workbook()
    ws = wb.active
    ws.title = "Historique des recherches"
    ws.append([
        # « Groupes » colle à « Utilisateur » : c'est sa qualification, pas
        # un critère de recherche. Les groupes sont ceux capturés AU MOMENT
        # de la recherche (search_log.log_search), pas l'appartenance
        # d'aujourd'hui — deux lignes du même utilisateur peuvent donc
        # légitimement différer s'il a changé de service.
        "Date / heure", "Utilisateur", "Groupes", "Requête", "Champ recherché",
        "Source(s)", "Extension(s)", "Auteur(s)", "Dossier",
        "Période début", "Période fin",
        "Résultats", "Documents retournés", "Avis", "Clics",
        # Ajoutées EN FIN de ligne et non à côté de « Résultats », où
        # elles auraient pourtant mieux leur place : décaler les colonnes
        # existantes casserait les classeurs et macros construits sur un
        # export précédent. Vides pour les recherches antérieures à la
        # mesure.
        "Durée (ms)", "Moteur ES (ms)",
    ])

    try:
        hits = es_scan(
            es,
            index=search_log.SEARCH_LOG_INDEX,
            query={"query": query, "sort": [{"timestamp": {"order": "desc"}}]},
            preserve_order=True,
        )
        for i, hit in enumerate(hits):
            if i >= SEARCH_LOGS_EXPORT_MAX_ROWS:
                break
            s = hit["_source"]
            ws.append([
                s.get("timestamp", ""),
                s.get("username", ""),
                _join(s.get("groups")),
                s.get("query", ""),
                s.get("search_in", ""),
                _join(s.get("source")),
                _join(s.get("extension")),
                _join(s.get("author")),
                _join(s.get("folder")),
                s.get("date_from", ""),
                s.get("date_to", ""),
                s.get("total_results", 0),
                _join(s.get("result_files")),
                s.get("feedback", ""),
                len(s.get("clicks") or []),
                s.get("duration_ms", ""),
                s.get("took_ms", ""),
            ])
    except Exception as e:
        if "index_not_found" not in str(e).lower():
            raise HTTPException(status_code=500, detail=str(e)) from e

    # Une largeur par colonne d'en-tête, dans le même ordre : les deux
    # listes doivent rester de même longueur.
    for col_idx, width in enumerate([19, 14, 28, 30, 14, 14, 14, 16, 20, 14, 14, 10, 40, 8, 8, 12, 14], start=1):
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = width

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    filename = f"historique-recherches-{datetime.now(timezone.utc).strftime('%Y%m%d-%H%M')}.xlsx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Pages ──────────────────────────────────────────────────────
# L'interface web (index.html, chat.html) est servie directement par
# Nginx depuis le projet docsearch-ui — cette API est maintenant une
# API JSON pure, sans dépendance sur des templates HTML.
# Voir docsearch-ui et la configuration nginx.conf de docsearch-infra.
